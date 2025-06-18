#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
AI校对助手使用示例
"""

import os
from proofreader import ProofReader


def example_quick_check():
    """快速检查文本示例"""
    print("=== 快速文本检查示例 ===")
    
    # 示例文本（包含一些常见问题）
    test_text = """
    计算器科学是一门研究算法和数据结构的学科。在程式设计中，我们经常使用变数来存储数据。
    Python是一种高级程序设计语言，它具有简洁、易读的语法。在编写代码时，我们需要注意语法的正确性。
    人工智能技术在计算机领域得到了广泛应用，特别是在机器学习和深度学习方面。
    """
    
    try:
        # 创建校对器
        proofreader = ProofReader()
        
        # 执行快速检查
        result = proofreader.quick_check(test_text)
        
        print(f"检查完成！发现 {len(result.issues)} 个问题")
        
        # 显示发现的问题
        if result.issues:
            print("\n发现的问题：")
            for i, issue in enumerate(result.issues, 1):
                print(f"{i}. {issue['type']}: {issue['text']}")
                print(f"   建议: {issue['suggestion']}")
                print(f"   严重程度: {issue['severity']}")
                print()
        
        # 显示改进建议
        if result.suggestions:
            print("改进建议：")
            for i, suggestion in enumerate(result.suggestions, 1):
                print(f"{i}. 原文: {suggestion['original']}")
                print(f"   建议: {suggestion['suggested']}")
                print(f"   原因: {suggestion['reason']}")
                print()
    
    except Exception as e:
        print(f"执行失败: {e}")
        print("请确保已正确配置 OpenAI API Key")


def example_document_proofrading():
    """文档校对示例"""
    print("\n=== 文档校对示例 ===")
    
    # 注意：这里需要真实的docx文件
    input_file = "example_input.docx"
    output_file = "example_output.docx"
    
    if not os.path.exists(input_file):
        print(f"示例输入文件 {input_file} 不存在")
        print("请创建一个名为 example_input.docx 的测试文件")
        return
    
    try:
        # 创建校对器
        proofreader = ProofReader()
        
        # 执行文档校对
        success = proofreader.proofread_document(input_file, output_file)
        
        if success:
            print(f"文档校对成功！输出文件: {output_file}")
        else:
            print("文档校对失败")
    
    except Exception as e:
        print(f"执行失败: {e}")


def example_batch_proofrading():
    """批量校对示例"""
    print("\n=== 批量校对示例 ===")
    
    input_dir = "./test_documents"
    output_dir = "./proofread_documents"
    
    if not os.path.exists(input_dir):
        print(f"输入目录 {input_dir} 不存在")
        os.makedirs(input_dir)
        print(f"已创建目录 {input_dir}，请在其中放入要校对的docx文件")
        return
    
    try:
        # 创建校对器
        proofreader = ProofReader()
        
        # 执行批量校对
        success = proofreader.batch_proofread(input_dir, output_dir)
        
        if success:
            print(f"批量校对成功！输出目录: {output_dir}")
        else:
            print("批量校对失败")
    
    except Exception as e:
        print(f"执行失败: {e}")


def create_sample_docx():
    """创建示例docx文件"""
    print("\n=== 创建示例文档 ===")
    
    try:
        from docx import Document
        
        # 创建新文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('计算机科学基础', 0)
        
        # 添加段落
        p1 = doc.add_paragraph('计算器科学是研究算法、数据结构和计算系统设计的学科。')
        p2 = doc.add_paragraph('在程式设计过程中，我们需要使用各种编程语言，如Python、Java等。')
        p3 = doc.add_paragraph('变数是程序中用来存储数据的容器，它们可以存储不同类型的数据。')
        p4 = doc.add_paragraph('函式是一段可重用的代码，它接受输入参数并返回输出结果。')
        p5 = doc.add_paragraph('人工智能技术在现代计算机科学中占据重要地位，包括机器学习、深度学习等领域。')
        
        # 保存文档
        filename = 'example_input.docx'
        doc.save(filename)
        print(f"已创建示例文档: {filename}")
        
        return filename
    
    except ImportError:
        print("需要安装 python-docx 库")
        return None
    except Exception as e:
        print(f"创建示例文档失败: {e}")
        return None


if __name__ == "__main__":
    print("AI校对助手使用示例")
    print("=" * 50)
    
    # 检查环境
    if not os.getenv('OPENAI_API_KEY'):
        print("⚠️  警告: 未设置 OPENAI_API_KEY 环境变量")
        print("请先设置 API Key 才能正常使用校对功能")
        print()
    
    # 运行示例
    example_quick_check()
    
    # 创建示例文档
    sample_file = create_sample_docx()
    
    if sample_file:
        example_document_proofrading()
    
    # 批量校对示例
    example_batch_proofrading()
    
    print("\n示例运行完成！")
    print("更多使用方法请参考 README.md 文档") 
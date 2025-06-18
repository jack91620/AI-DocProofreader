#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
测试真正的Word跟踪更改功能
"""

from docx import Document
from proofreader.word_track_changes import WordTrackChangesManager, enable_track_changes_in_docx
import os
import zipfile
import tempfile

def create_test_document():
    """创建测试文档"""
    doc = Document()
    doc.add_heading('测试文档 - Word跟踪更改功能', 0)
    
    doc.add_paragraph("这是一个测试段落。")
    doc.add_paragraph("计算器科学是一门非常重要的学科，涉及到程式设计和筭法等内容。")
    doc.add_paragraph("在日常生活中，我们经常需要进行文字校对工作。")
    doc.add_paragraph("AI技术的发展为文字处理带来了新的可能性。")
    
    return doc

def test_real_word_track_changes():
    """测试真正的Word跟踪更改功能"""
    print("🔄 开始测试真正的Word跟踪更改功能...")
    
    # 创建测试文档
    doc = create_test_document()
    print("✅ 创建测试文档")
    
    # 创建跟踪更改管理器
    track_manager = WordTrackChangesManager(doc)
    print("✅ 创建跟踪更改管理器")
    
    # 添加跟踪更改
    paragraphs = list(doc.paragraphs)
    
    changes = [
        (paragraphs[2], "计算器科学", "计算机科学", "错别字修正"),
        (paragraphs[2], "程式设计", "程序设计", "术语统一"),
        (paragraphs[2], "筭法", "算法", "错别字修正"),
    ]
    
    for paragraph, original, corrected, reason in changes:
        success = track_manager.add_tracked_change(paragraph, original, corrected, reason)
        if success:
            print(f"✅ 添加跟踪更改: {original} -> {corrected}")
        else:
            print(f"❌ 添加跟踪更改失败: {original}")
    
    # 应用所有跟踪更改
    print("✅ 应用所有跟踪更改到文档")
    track_manager.apply_all_changes()
    
    # 保存临时文档
    temp_file = "test_track_changes_temp.docx"
    doc.save(temp_file)
    print("✅ 保存临时文档")
    
    # 启用跟踪更改并生成最终文档
    output_file = "test_real_track_changes.docx"
    success = enable_track_changes_in_docx(temp_file, output_file, track_manager.revisions_data)
    
    if success:
        print(f"✅ 成功生成带有真正Word跟踪更改的文档: {output_file}")
        
        # 验证XML结构
        verify_track_changes_xml(output_file)
        
        # 清理临时文件
        if os.path.exists(temp_file):
            os.remove(temp_file)
            
        print("\n📝 使用方法:")
        print("1. 用Microsoft Word打开生成的文档")
        print("2. 在'审阅'选项卡中查看跟踪更改")
        print("3. 可以接受或拒绝每个修改")
        
    else:
        print("❌ 生成失败")
    
    return success

def verify_track_changes_xml(docx_file):
    """验证跟踪更改的XML结构"""
    print("\n🔍 验证Word跟踪更改XML结构...")
    
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            # 解压docx文件
            with zipfile.ZipFile(docx_file, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # 检查document.xml
            document_path = os.path.join(temp_dir, 'word', 'document.xml')
            if os.path.exists(document_path):
                with open(document_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # 统计修订标记
                del_count = content.count('<w:del ')
                ins_count = content.count('<w:ins ')
                deltext_count = content.count('<w:delText>')
                
                print(f"📊 document.xml修订标记统计:")
                print(f"   - w:del (删除标记): {del_count}")
                print(f"   - w:ins (插入标记): {ins_count}")
                print(f"   - w:delText (删除文本): {deltext_count}")
                
                if del_count > 0 and ins_count > 0:
                    print("✅ 发现真正的Word修订标记")
                else:
                    print("⚠️  未发现Word修订标记")
            
            # 检查settings.xml
            settings_path = os.path.join(temp_dir, 'word', 'settings.xml')
            if os.path.exists(settings_path):
                with open(settings_path, 'r', encoding='utf-8') as f:
                    settings_content = f.read()
                
                if 'trackRevisions' in settings_content:
                    print("✅ settings.xml中已启用跟踪更改")
                else:
                    print("⚠️  settings.xml中未启用跟踪更改")
            
    except Exception as e:
        print(f"❌ 验证XML结构失败: {e}")

if __name__ == "__main__":
    test_real_word_track_changes() 
#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
创建包含各种问题的示例docx文档
"""

from docx import Document
from docx.shared import Inches

def create_sample_document():
    """创建包含各种校对问题的示例文档"""
    
    doc = Document()
    
    # 添加标题
    doc.add_heading('计算机科学基础教材', 0)
    
    # 第一章 - 包含错别字问题
    doc.add_heading('第一章 计算器科学概述', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run('计算器科学是一门研究算法、数据结构和计算系统设计的学科。')
    p1.add_run('它涉及到程式设计、软体工程、人工智能等多个领域。')
    
    p2 = doc.add_paragraph()
    p2.add_run('在现代社会中，计算机技术已经成为推动科技发展的重要力量。')
    p2.add_run('从个人电脑到超级计算器，从移动设备到云端服务，')
    p2.add_run('计算机无处不在地影响着我们的生活。')
    
    # 第二章 - 包含术语不一致问题
    doc.add_heading('第二章 程序设计基础', level=1)
    
    p3 = doc.add_paragraph()
    p3.add_run('程式设计是计算机科学的核心技能之一。')
    p3.add_run('在程序开发过程中，我们需要掌握各种编程语言和程式设计技巧。')
    
    p4 = doc.add_paragraph()
    p4.add_run('变数是程序中用来存储数据的容器。')
    p4.add_run('在Python中，变量可以存储不同类型的数据，')
    p4.add_run('如整数、浮点数、字符串等。')
    
    p5 = doc.add_paragraph()
    p5.add_run('函式是一段可重用的代码块，它可以接受参数并返回结果。')
    p5.add_run('通过使用函数，我们可以让代码更加模块化和可维护。')
    
    # 第三章 - 包含语法和标点问题
    doc.add_heading('第三章 数据结构与算法', level=1)
    
    p6 = doc.add_paragraph()
    p6.add_run('数据结构是组织和存储数据的方式,常见的数据结构包括数组、链表、栈、队列等。')
    p6.add_run('选择合适的数据结构对程序的性能和效率至关重要。')
    
    p7 = doc.add_paragraph()
    p7.add_run('算法是解决问题的步骤和方法。一个好的算法应该具备正确性、有效性和可读性。')
    p7.add_run('时间复杂度和空间复杂度是评估算法性能的重要指标')  # 缺少标点
    
    p8 = doc.add_paragraph()
    p8.add_run('排序算法是计算机科学中最基本的算法之一。')
    p8.add_run('常见的排序算法有冒泡排序，选择排序，插入排序，快速排序等。')
    p8.add_run('不同的排序算法在不同情况下有着不同的性能表现，')
    
    # 第四章 - 包含专业术语问题
    doc.add_heading('第四章 软件工程', level=1)
    
    p9 = doc.add_paragraph()
    p9.add_run('软体工程是一门研究如何系统化、规范化、量化地开发和维护软件的学科。')
    p9.add_run('它涉及软件开发的整个生命周期，从需求分析到系统维护。')
    
    p10 = doc.add_paragraph()
    p10.add_run('在软件开发过程中，我们需要使用各种工具和方法论。')
    p10.add_run('敏捷开发、瀑布模型、螺旋模型等都是常用的软件开发方法。')
    
    p11 = doc.add_paragraph()
    p11.add_run('版本控制系统如Git可以帮助我们管理代码变更，')
    p11.add_run('测试框架可以确保软件质量，')
    p11.add_run('持续集成和持续部署(CI/CD)可以提高开发效率。')
    
    # 第五章 - 包含更多问题
    doc.add_heading('第五章 人工智能与机器学习', level=1)
    
    p12 = doc.add_paragraph()
    p12.add_run('人工智能（AI）是计算机科学的一个重要分支，')
    p12.add_run('它致力于创建能够模拟人类智能的计算机系统。')
    p12.add_run('机器学习是实现人工智能的重要手段之一')  # 缺少标点
    
    p13 = doc.add_paragraph()
    p13.add_run('深度学习作为机器学习的子领域，通过多层神经网络来学习数据的复杂模式。')
    p13.add_run('卷积神经网络(CNN)在图像识别领域表现出色，')
    p13.add_run('循环神经网络(RNN)在自然语言处理任务中很有用。')
    
    p14 = doc.add_paragraph()
    p14.add_run('自然语言处理(NLP)技术使得计算机能够理解和生成人类语言，')
    p14.add_run('这为智能聊天机器人、机器翻译、文本摘要等应用提供了基础。')
    
    # 保存文档
    filename = 'sample_input.docx'
    doc.save(filename)
    print(f"✅ 已创建示例文档: {filename}")
    
    # 显示文档统计信息
    paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
    total_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    char_count = len(total_text)
    
    print(f"📊 文档统计：")
    print(f"  - 段落数：{paragraph_count}")
    print(f"  - 字符数：{char_count}")
    print(f"  - 包含的问题类型：")
    print(f"    • 错别字：计算器科学 → 计算机科学")
    print(f"    • 术语不一致：程式设计 vs 程序设计")
    print(f"    • 标点符号：缺少句号")
    print(f"    • 专业术语：软体工程 → 软件工程")
    
    return filename

if __name__ == "__main__":
    create_sample_document() 
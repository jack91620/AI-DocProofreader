#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
展示Word批注效果的脚本
"""

from docx import Document
from rich.console import Console
from rich.table import Table
from rich.panel import Panel

console = Console()

def analyze_word_comments(file_path):
    """分析文档中的Word批注"""
    try:
        doc = Document(file_path)
        
        console.print(f"[bold blue]📝 分析文档: {file_path}[/bold blue]")
        
        # 统计信息
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        
        # 查找批注标记
        comment_indicators = 0
        highlighted_text = 0
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 检查是否有批注标记（💬符号）
                if '💬' in paragraph.text:
                    comment_indicators += 1
                
                # 检查run级别的格式
                for run in paragraph.runs:
                    # 检查是否有高亮
                    if run.font.highlight_color is not None:
                        highlighted_text += 1
        
        # 显示统计
        table = Table(title="Word批注分析结果")
        table.add_column("项目", style="cyan")
        table.add_column("数量", style="green")
        
        table.add_row("总段落数", str(paragraph_count))
        table.add_row("批注标记(💬)", str(comment_indicators))
        table.add_row("高亮文本runs", str(highlighted_text))
        
        console.print(table)
        
        # 显示带批注的段落
        console.print(f"\n[bold yellow]📋 带有批注标记的段落：[/bold yellow]")
        
        for i, paragraph in enumerate(doc.paragraphs, 1):
            if paragraph.text.strip() and '💬' in paragraph.text:
                console.print(f"[cyan]段落 {i}:[/cyan] {paragraph.text}")
        
        return comment_indicators > 0
        
    except Exception as e:
        console.print(f"[red]分析失败: {e}[/red]")
        return False

def compare_comment_systems():
    """对比不同的批注系统效果"""
    console.print(Panel.fit("[bold blue]Word批注系统对比分析[/bold blue]"))
    
    files_to_analyze = [
        ("sample_input.docx", "原始输入文档"),
        ("sample_output_with_comments.docx", "旧版批注系统"),
        ("sample_output_with_word_comments.docx", "新版Word批注系统")
    ]
    
    for filename, description in files_to_analyze:
        console.print(f"\n{'='*60}")
        console.print(f"[bold green]{description}[/bold green]")
        
        try:
            has_comments = analyze_word_comments(filename)
            
            if filename == "sample_output_with_word_comments.docx" and has_comments:
                console.print("[green]✅ 检测到Word批注标记！[/green]")
            elif filename == "sample_input.docx":
                console.print("[blue]ℹ️ 这是原始文档，无批注[/blue]")
            else:
                console.print("[yellow]⚠️ 未检测到明显的批注标记[/yellow]")
                
        except FileNotFoundError:
            console.print(f"[red]❌ 文件不存在: {filename}[/red]")

def show_batch_comment_details():
    """显示批注功能的详细信息"""
    console.print(f"\n[bold blue]🎯 Word批注功能说明[/bold blue]")
    
    features = [
        "✅ 使用Word原生XML批注结构",
        "✅ 高亮显示有问题的文本（黄色背景）",
        "✅ 添加批注标记（💬 符号）",
        "✅ 在XML级别添加批注范围标记",
        "✅ 尝试创建comments.xml部分",
        "✅ 提供详细的批注信息记录",
        "✅ 多级备用机制确保功能稳定",
    ]
    
    for feature in features:
        console.print(f"  {feature}")
    
    console.print(f"\n[bold cyan]📖 使用方式：[/bold cyan]")
    console.print("1. 用Microsoft Word打开生成的文档")
    console.print("2. 查看文本中的高亮部分")
    console.print("3. 注意💬符号标记的批注位置")
    console.print("4. 在Word中检查是否有审阅批注")
    
    console.print(f"\n[bold yellow]🔧 技术细节：[/bold yellow]")
    console.print("• 高亮显示：使用WD_COLOR_INDEX.YELLOW")
    console.print("• 批注标记：插入💬符号作为视觉指示")
    console.print("• XML结构：添加commentRangeStart/End标记")
    console.print("• 备用机制：确保在各种情况下都能正常工作")

if __name__ == "__main__":
    compare_comment_systems()
    show_batch_comment_details()
    
    console.print(f"\n[bold green]🎉 Word批注系统升级完成！[/bold green]")
    console.print("现在系统使用更接近Microsoft Word原生批注的方式来标记问题。")
    console.print("虽然python-docx对完整批注功能的支持有限，但我们通过以下方式改进了用户体验：")
    console.print("• 高亮显示问题文本")
    console.print("• 添加视觉批注标记")
    console.print("• 在XML级别添加批注结构") 
    console.print("• 提供详细的校对报告") 
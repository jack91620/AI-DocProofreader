#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
测试带批注的真正Word跟踪更改功能
"""

from docx import Document
from proofreader.word_track_changes_with_comments import WordTrackChangesWithCommentsManager, enable_track_changes_and_comments_in_docx
import os


def test_track_changes_with_comments():
    """测试带批注的跟踪更改功能"""
    try:
        print("🔄 开始测试带批注的跟踪更改功能...")
        
        # 创建测试文档
        doc = Document()
        doc.add_heading('测试文档 - 带批注的跟踪更改', 0)
        doc.add_paragraph("这是一个测试段落。")
        doc.add_paragraph("计算器科学是一门非常重要的学科，涉及到程式设计和筭法等内容。")
        doc.add_paragraph("在日常生活中，我们经常需要进行文字校对工作。")
        doc.add_paragraph("AI技术的发展为文字处理带来了新的可能性。")
        
        # 创建带批注的跟踪更改管理器
        track_comments_manager = WordTrackChangesWithCommentsManager(doc)
        
        # 添加带批注的跟踪更改
        paragraphs = list(doc.paragraphs)
        
        changes = [
            (paragraphs[2], "计算器科学", "计算机科学", "错别字修正：'器'应为'机'"),
            (paragraphs[2], "程式设计", "程序设计", "术语统一：使用标准中文术语"),
            (paragraphs[2], "筭法", "算法", "错别字修正：'筭'应为'算'"),
        ]
        
        print(f"准备修改 {len(changes)} 个问题...")
        
        for paragraph, original, corrected, reason in changes:
            success = track_comments_manager.add_tracked_change_with_comment(
                paragraph, original, corrected, reason
            )
            if not success:
                print(f"❌ 修改失败: {original} -> {corrected}")
        
        # 应用所有更改
        track_comments_manager.apply_all_changes()
        
        # 获取统计信息
        stats = track_comments_manager.get_statistics()
        print(f"\n📊 修订统计:")
        print(f"   - 总修改数: {stats['total_changes']}")
        print(f"   - 跟踪更改数: {stats['track_changes_count']}")
        print(f"   - 批注数: {stats['comments_count']}")
        print(f"   - 成功率: {stats['success_rate']:.1f}%")
        
        # 保存临时文档
        temp_file = "test_track_changes_with_comments_temp.docx"
        doc.save(temp_file)
        print(f"✅ 保存临时文档: {temp_file}")
        
        # 生成最终文档
        output_file = "test_track_changes_with_comments.docx"
        success = enable_track_changes_and_comments_in_docx(
            temp_file, 
            output_file, 
            track_comments_manager.track_changes_manager.revisions_data,
            track_comments_manager.comments_manager.comments
        )
        
        if success:
            print(f"\n✅ 带批注的跟踪更改文档已创建: {output_file}")
            print("📝 在Microsoft Word中可以看到:")
            print("   - 🔄 真正的跟踪更改（红色删除线 + 蓝色下划线）")
            print("   - 💬 详细的批注说明（修订原因和类型）")
            print("   - ✅ 可以接受/拒绝修改")
            print("   - 💭 可以回复批注")
            
            # 清理临时文件
            if os.path.exists(temp_file):
                os.remove(temp_file)
                print(f"🗑️ 清理临时文件: {temp_file}")
        else:
            print("❌ 创建失败")
        
        return success
        
    except Exception as e:
        print(f"❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def verify_output_file():
    """验证输出文件"""
    output_file = "test_track_changes_with_comments.docx"
    
    if not os.path.exists(output_file):
        print(f"❌ 输出文件不存在: {output_file}")
        return False
    
    try:
        # 尝试用python-docx读取文件验证
        doc = Document(output_file)
        print(f"✅ 文件验证通过: {len(doc.paragraphs)} 个段落")
        
        # 调用验证脚本
        from verify_real_track_changes import verify_word_track_changes
        return verify_word_track_changes(output_file)
        
    except Exception as e:
        print(f"❌ 文件验证失败: {e}")
        return False


if __name__ == "__main__":
    print("🚀 开始测试带批注的Word跟踪更改功能")
    print("=" * 60)
    
    # 运行测试
    success = test_track_changes_with_comments()
    
    if success:
        print("\n" + "=" * 60)
        print("🔍 验证输出文件...")
        verify_output_file()
        
        print("\n" + "=" * 60)
        print("✅ 测试完成！")
        print("📁 请用Microsoft Word查看生成的文档：")
        print("   - test_track_changes_with_comments.docx")
    else:
        print("\n❌ 测试失败！") 
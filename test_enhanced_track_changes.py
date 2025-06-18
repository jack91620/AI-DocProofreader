#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
测试增强版跟踪更改功能 - 同时使用跟踪更改和批注
"""

from docx import Document
from proofreader.word_track_changes import WordTrackChangesManager, enable_track_changes_in_docx
from proofreader.word_comments_advanced import WordCommentsManager
import os


def create_test_document():
    """创建测试文档"""
    doc = Document()
    doc.add_heading('AI校对测试文档', 0)
    
    doc.add_paragraph("这是一个测试段落。")
    doc.add_paragraph("计算器科学是一门非常重要的学科，涉及到程式设计和筭法等内容。")
    doc.add_paragraph("在日常生活中，我们经常需要进行文字校对工作。")
    doc.add_paragraph("AI技术的发展为文字处理带来了新的可能性。")
    
    return doc


def add_track_changes_with_comments(doc):
    """添加跟踪更改和批注"""
    print("🔄 开始添加跟踪更改和批注...")
    
    # 获取段落
    paragraphs = list(doc.paragraphs)
    target_paragraph = paragraphs[2]  # 第三个段落
    
    # 创建跟踪更改管理器
    track_manager = WordTrackChangesManager(doc)
    
    # 修订数据
    revisions = [
        ("计算器科学", "计算机科学", "错别字修正：'器'应为'机'"),
        ("程式设计", "程序设计", "术语统一：使用标准中文术语"),
        ("筭法", "算法", "错别字修正：'筭'应为'算'"),
    ]
    
    # 先添加跟踪更改
    print("📝 添加跟踪更改...")
    for original, corrected, reason in revisions:
        success = track_manager.add_tracked_change(target_paragraph, original, corrected, reason)
        if success:
            print(f"✅ 跟踪更改: {original} -> {corrected}")
        else:
            print(f"❌ 跟踪更改失败: {original}")
    
    # 应用所有跟踪更改
    print("🔧 应用跟踪更改...")
    track_manager.apply_all_changes()
    
    return track_manager


def create_enhanced_document():
    """创建增强版文档（跟踪更改+批注）"""
    try:
        print("🚀 开始创建增强版文档...")
        
        # 创建测试文档
        doc = create_test_document()
        print("✅ 创建基础文档")
        
        # 添加跟踪更改
        track_manager = add_track_changes_with_comments(doc)
        
        # 保存临时文档
        temp_file = "test_enhanced_temp.docx"
        doc.save(temp_file)
        print(f"✅ 保存临时文档: {temp_file}")
        
        # 生成最终文档
        output_file = "test_enhanced_track_changes_comments.docx"
        success = enable_track_changes_in_docx(temp_file, output_file, track_manager.revisions_data)
        
        if success:
            print(f"✅ 增强版文档创建成功: {output_file}")
            
            # 现在添加批注到最终文档
            print("💬 添加批注...")
            add_comments_to_final_document(output_file)
            
            # 清理临时文件
            if os.path.exists(temp_file):
                os.remove(temp_file)
                print(f"🗑️ 清理临时文件: {temp_file}")
            
            return True
        else:
            print("❌ 文档创建失败")
            return False
            
    except Exception as e:
        print(f"❌ 创建增强版文档失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def add_comments_to_final_document(docx_file):
    """为最终文档添加批注"""
    try:
        # 重新读取文档
        doc = Document(docx_file)
        
        # 创建批注管理器
        comments_manager = WordCommentsManager(doc)
        
        # 批注数据
        comment_data = [
            ("计算机科学", "🔄 修订说明：\n原文：'计算器科学'\n修正：'计算机科学'\n原因：错别字修正，'器'应为'机'\n类型：错别字纠正"),
            ("程序设计", "🔄 修订说明：\n原文：'程式设计'\n修正：'程序设计'\n原因：术语统一，使用标准中文术语\n类型：术语规范化"),
            ("算法", "🔄 修订说明：\n原文：'筭法'\n修正：'算法'\n原因：错别字修正，'筭'应为'算'\n类型：错别字纠正"),
        ]
        
        # 添加批注
        comment_count = 0
        for text, comment in comment_data:
            # 在第二个段落（索引为2）中查找文本
            if len(doc.paragraphs) > 2:
                paragraph = doc.paragraphs[2]
                if comments_manager.add_comment(paragraph, text, comment):
                    comment_count += 1
                    print(f"✅ 添加批注 {comment_count}: {text}")
                else:
                    print(f"⚠️ 批注添加失败: {text}")
        
        # 保存文档
        enhanced_output = docx_file.replace('.docx', '_with_comments.docx')
        doc.save(enhanced_output)
        print(f"✅ 带批注的文档已保存: {enhanced_output}")
        
        return comment_count > 0
        
    except Exception as e:
        print(f"❌ 添加批注失败: {e}")
        return False


def verify_enhanced_document():
    """验证增强版文档"""
    files_to_check = [
        "test_enhanced_track_changes_comments.docx",
        "test_enhanced_track_changes_comments_with_comments.docx"
    ]
    
    print("\n🔍 验证增强版文档...")
    
    for filename in files_to_check:
        if os.path.exists(filename):
            print(f"\n📄 检查文件: {filename}")
            
            # 调用验证脚本
            try:
                from verify_real_track_changes import verify_word_track_changes
                success = verify_word_track_changes(filename)
                if success:
                    print(f"✅ {filename} 验证通过")
                else:
                    print(f"⚠️ {filename} 验证有问题")
            except Exception as e:
                print(f"❌ 验证失败: {e}")
        else:
            print(f"⚠️ 文件不存在: {filename}")


if __name__ == "__main__":
    print("🌟 测试增强版Word跟踪更改+批注功能")
    print("=" * 60)
    
    # 创建增强版文档
    success = create_enhanced_document()
    
    if success:
        print("\n" + "=" * 60)
        print("🔍 验证生成的文档...")
        verify_enhanced_document()
        
        print("\n" + "=" * 60)
        print("✅ 测试完成！")
        print("📁 生成的文档：")
        print("   - test_enhanced_track_changes_comments.docx (跟踪更改版)")
        print("   - test_enhanced_track_changes_comments_with_comments.docx (跟踪更改+批注版)")
        print("\n📝 在Microsoft Word中可以看到：")
        print("   - 🔄 真正的跟踪更改（红色删除线 + 蓝色下划线）")
        print("   - 💬 详细的批注说明（修订原因和类型）")
        print("   - ✅ 可以接受/拒绝修改")
        print("   - 💭 可以回复批注")
    else:
        print("\n❌ 测试失败！") 
#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
修复版增强校对器 - 确保跟踪更改和批注都正确显示
"""

import os
import sys
from typing import Optional
from rich.console import Console
from docx import Document
from datetime import datetime
import zipfile
import tempfile
import xml.etree.ElementTree as ET

from .config import Config
from .document import DocumentProcessor
from .ai_checker import AIChecker, ProofreadingResult
from .word_track_changes import WordTrackChangesManager, enable_track_changes_in_docx
from .word_comments_advanced import WordCommentsManager
from .word_comments_xml import create_comments_xml, create_document_rels, update_content_types


class ProofReaderWithTrackChangesAndCommentsFixed:
    """修复版增强校对器 - 确保跟踪更改和批注都正确显示"""
    
    def __init__(self, api_key: str = None):
        """初始化校对器"""
        if api_key:
            # 如果传入了API密钥，设置环境变量
            import os
            os.environ['OPENAI_API_KEY'] = api_key
        self.config = Config()
        self.ai_checker = AIChecker(self.config)
        self.console = Console()
        self.document_processor = DocumentProcessor()
    
    def proofread_with_track_changes_and_comments(self, input_file: str, output_file: str = None) -> bool:
        """使用跟踪更改和批注进行校对 - 修复版"""
        try:
            # 生成输出文件名
            if not output_file:
                output_file = input_file.replace('.docx', '_enhanced_fixed.docx')
            
            self.console.print(f"[green]开始修复版增强校对：{input_file}[/green]")
            
            # 第一步：AI校对
            self.console.print("[blue]第一步：AI校对分析文档...[/blue]")
            doc = Document(input_file)
            text_content = self.extract_text_content(doc)
            self.console.print(f"[blue]提取文本内容: {len(text_content)} 个段落[/blue]")
            
            ai_result = self.ai_checker.check_text(' '.join(text_content))
            
            # 第二步：创建同步更改数据
            self.console.print("[blue]第二步：生成同步更改数据...[/blue]")
            synchronized_changes = self._create_synchronized_changes(ai_result, text_content, doc)
            self.console.print(f"[green]✅ 发现 {len(synchronized_changes)} 个需要修改的问题[/green]")
            
            # 第三步：应用更改和批注
            self.console.print("[blue]第三步：应用跟踪更改和批注...[/blue]")
            success = self._apply_changes_with_proper_comments(doc, synchronized_changes, output_file)
            
            if success:
                self.console.print(f"[green]✅ 修复版增强校对完成：{output_file}[/green]")
                self.console.print("[blue]📝 文档包含：[/blue]")
                self.console.print("   - 🔄 Word跟踪更改（可接受/拒绝）")
                self.console.print("   - 💬 对应的详细批注（可查看/回复）")
                self.console.print("   - 🔗 正确的批注引用链接")
                return True
            else:
                return False
            
        except Exception as e:
            self.console.print(f"[red]❌ 修复版增强校对失败: {e}[/red]")
            import traceback
            traceback.print_exc()
            return False

    def _create_synchronized_changes(self, ai_result: ProofreadingResult, text_content: list, doc: Document):
        """创建同步的跟踪更改和批注数据"""
        synchronized_changes = []
        processed_pairs = set()  # 避免重复处理相同的修正对
        
        self.console.print(f"[blue]🔍 处理AI发现的 {len(ai_result.issues)} 个问题和 {len(ai_result.suggestions)} 个建议[/blue]")
        
        # 创建所有可能的修正对
        all_corrections = []
        
        # 从suggestions中提取修正对
        for suggestion in ai_result.suggestions:
            original_text = suggestion.get('original', '')
            suggested_text = suggestion.get('suggested', '')
            reason = suggestion.get('reason', '')
            
            if suggested_text and suggested_text != original_text:
                # 提取具体的词汇修正
                corrections = self._extract_word_corrections(original_text, suggested_text)
                for orig, corr in corrections:
                    correction_pair = (orig, corr)
                    if correction_pair not in processed_pairs:
                        all_corrections.append({
                            'original': orig,
                            'corrected': corr,
                            'reason': reason,
                            'type': 'suggestion',
                            'full_original': original_text,
                            'full_suggested': suggested_text
                        })
                        processed_pairs.add(correction_pair)
        
        # 从issues中提取修正对
        for issue in ai_result.issues:
            problem_text = issue.get('text', '')
            suggestion = issue.get('suggestion', '')
            issue_type = issue.get('type', '')
            severity = issue.get('severity', '')
            
            if issue_type == "术语不一致" and "发现多种术语：" in problem_text:
                # 处理术语不一致
                terms = self._extract_terms_from_inconsistency(problem_text, suggestion)
                for original_term, corrected_term in terms:
                    correction_pair = (original_term, corrected_term)
                    if correction_pair not in processed_pairs:
                        all_corrections.append({
                            'original': original_term,
                            'corrected': corrected_term,
                            'reason': f"{issue_type} - {severity}",
                            'type': 'terminology',
                            'suggestion_text': suggestion
                        })
                        processed_pairs.add(correction_pair)
            elif issue_type in ["错别字和用词不当", "标点符号使用"]:
                # 处理错别字和标点问题
                corrected_text = self._extract_corrected_text(suggestion)
                if corrected_text and corrected_text != problem_text:
                    correction_pair = (problem_text, corrected_text)
                    if correction_pair not in processed_pairs:
                        all_corrections.append({
                            'original': problem_text,
                            'corrected': corrected_text,
                            'reason': f"{issue_type} - {severity}",
                            'type': 'error_fix',
                            'suggestion_text': suggestion
                        })
                        processed_pairs.add(correction_pair)
        
        # 创建段落索引映射 - 从非空段落索引到实际段落索引
        paragraph_mapping = {}
        text_index = 0
        for doc_index, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                paragraph_mapping[text_index] = doc_index
                text_index += 1
        
        # 改进的文本匹配和应用逻辑
        for correction in all_corrections:
            original = correction['original']
            corrected = correction['corrected']
            reason = correction['reason']
            corr_type = correction['type']
            
            # 在所有段落中查找匹配项
            matches_found = []
            for i, paragraph_text in enumerate(text_content):
                # 使用更精确的匹配策略
                if self._is_text_match(original, paragraph_text):
                    actual_paragraph_index = paragraph_mapping.get(i, i)
                    # 计算该术语在段落中出现的次数
                    occurrences = paragraph_text.count(original)
                    matches_found.append((i, actual_paragraph_index, paragraph_text, occurrences))
            
            # 处理所有匹配项
            if matches_found:
                for text_idx, para_idx, para_text, occurrences in matches_found:
                    # 为每个出现的术语创建一个修正项
                    for occurrence in range(occurrences):
                        # 创建批注文本
                        if corr_type == 'suggestion':
                            comment_text = f"💡 改进建议: {original} → {corrected}\n"
                            comment_text += f"📋 原因: {reason}\n"
                            comment_text += f"🎯 类型: 改进建议\n"
                        elif corr_type == 'terminology':
                            comment_text = f"🔍 术语不一致修正: {original} → {corrected}\n"
                            comment_text += f"📝 理由: {reason}\n"
                            comment_text += f"💡 建议: {correction.get('suggestion_text', '')}\n"
                        else:
                            comment_text = f"🔧 错误修正: {original} → {corrected}\n"
                            comment_text += f"📝 理由: {reason}\n"
                            comment_text += f"💡 建议: {correction.get('suggestion_text', '')}\n"
                        
                        comment_text += f"⏰ 处理时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                        
                        synchronized_changes.append({
                            'paragraph_index': para_idx,
                            'original_text': original,
                            'corrected_text': corrected,
                            'comment_text': comment_text,
                            'reason': reason,
                            'type': corr_type,
                            'occurrence_index': occurrence  # 添加出现次数索引
                        })
                        
                        self.console.print(f"[green]✅ 添加修正: {original} → {corrected} (段落{para_idx+1}, 第{occurrence+1}次出现)[/green]")
            else:
                self.console.print(f"[yellow]⚠️  未找到匹配文本: {original}[/yellow]")
        
        self.console.print(f"[green]✅ 总共创建了 {len(synchronized_changes)} 个同步更改[/green]")
        return synchronized_changes

    def _is_text_match(self, target_text: str, paragraph_text: str) -> bool:
        """改进的文本匹配逻辑"""
        # 精确匹配
        if target_text in paragraph_text:
            return True
        
        # 去除标点符号后匹配
        import re
        target_clean = re.sub(r'[^\w\s]', '', target_text)
        paragraph_clean = re.sub(r'[^\w\s]', '', paragraph_text)
        if target_clean in paragraph_clean:
            return True
        
        # 分词匹配（处理术语）
        target_words = target_text.split()
        if len(target_words) == 1 and target_words[0] in paragraph_text:
            return True
        
        return False

    def _apply_changes_with_proper_comments(self, doc: Document, synchronized_changes: list, output_file: str) -> bool:
        """应用更改并确保批注正确显示"""
        try:
            # 创建跟踪更改管理器
            track_changes_manager = WordTrackChangesManager(doc)
            
            # 创建批注管理器
            comments_manager = WordCommentsManager(doc)
            
            # 应用每个更改
            applied_count = 0
            for change in synchronized_changes:
                paragraph_index = change.get('paragraph_index', 0)
                original_text = change.get('original_text', '')
                corrected_text = change.get('corrected_text', '')
                comment_text = change.get('comment_text', '')
                reason = change.get('reason', '')
                
                if paragraph_index < len(doc.paragraphs):
                    paragraph = doc.paragraphs[paragraph_index]
                    
                    # 应用跟踪更改
                    track_success = track_changes_manager.add_tracked_change(
                        paragraph, original_text, corrected_text, reason
                    )
                    
                    # 应用批注
                    comment_success = comments_manager.add_comment(
                        paragraph, original_text, comment_text
                    )
                    
                    if track_success and comment_success:
                        applied_count += 1
                        self.console.print(f"[green]✅ 应用更改 {applied_count}: {original_text} -> {corrected_text}[/green]")
            
            # 完成处理
            track_changes_manager.apply_all_changes()
            comments_manager.finalize_document()
            
            # 保存带有基本更改的文档
            temp_file = output_file.replace('.docx', '_temp.docx')
            doc.save(temp_file)
            
            # 创建完整的批注系统
            success = self._create_complete_comment_system(
                temp_file, 
                output_file, 
                self._prepare_comments_with_changes(comments_manager.get_comments_for_xml(), synchronized_changes)
            )
            
            # 清理临时文件
            if os.path.exists(temp_file):
                os.remove(temp_file)
            
            if success:
                self.console.print(f"[green]✅ 成功应用 {applied_count} 个更改和批注[/green]")
                return True
            else:
                return False
                
        except Exception as e:
            self.console.print(f"[red]❌ 应用更改失败: {e}[/red]")
            import traceback
            traceback.print_exc()
            return False

    def _create_complete_comment_system(self, temp_file: str, output_file: str, comments_data: list) -> bool:
        """创建完整的批注系统"""
        try:
            self.console.print(f"[cyan]🔧 创建完整的批注系统，包含 {len(comments_data)} 个批注[/cyan]")
            
            with zipfile.ZipFile(temp_file, 'r') as input_zip:
                with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as output_zip:
                    # 复制原有文件，但需要修改document.xml
                    for item in input_zip.infolist():
                        if item.filename == 'word/document.xml':
                            # 修改document.xml添加批注引用
                            document_xml = input_zip.read(item.filename).decode('utf-8')
                            modified_document_xml = self._add_comment_references_to_document(document_xml, comments_data)
                            output_zip.writestr(item.filename, modified_document_xml.encode('utf-8'))
                        elif item.filename not in ['word/comments.xml', 'word/_rels/document.xml.rels', '[Content_Types].xml']:
                            output_zip.writestr(item, input_zip.read(item.filename))
                    
                    # 创建批注XML
                    comments_xml = self._create_comments_xml(comments_data)
                    output_zip.writestr('word/comments.xml', comments_xml)
                    
                    # 更新关系文件
                    rels_xml = self._create_updated_rels(input_zip)
                    output_zip.writestr('word/_rels/document.xml.rels', rels_xml)
                    
                    # 更新内容类型
                    content_types_xml = self._create_updated_content_types(input_zip)
                    output_zip.writestr('[Content_Types].xml', content_types_xml)
            
            self.console.print("[green]✅ 完整的批注系统创建成功[/green]")
            return True
            
        except Exception as e:
            self.console.print(f"[red]❌ 创建完整批注系统失败: {e}[/red]")
            import traceback
            traceback.print_exc()
            return False

    def _create_comments_xml(self, comments_data: list) -> str:
        """创建批注XML内容"""
        xml_content = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'''
        
        for comment in comments_data:
            comment_id = comment.get('id', 1)
            author = comment.get('author', 'AI校对助手')
            date = comment.get('date', datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"))
            text = comment.get('text', '')
            
            xml_content += f'''
    <w:comment w:id="{comment_id}" w:author="{author}" w:date="{date}" w:initials="AI">
        <w:p>
            <w:r>
                <w:t>{text}</w:t>
            </w:r>
        </w:p>
    </w:comment>'''
        
        xml_content += '\n</w:comments>'
        return xml_content

    def _create_updated_rels(self, input_zip) -> str:
        """创建更新的关系文件"""
        try:
            # 读取原始关系文件
            if 'word/_rels/document.xml.rels' in input_zip.namelist():
                rels_xml = input_zip.read('word/_rels/document.xml.rels').decode('utf-8')
            else:
                # 创建基本的关系文件
                rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>'''
            
            # 检查是否已包含批注关系
            if 'comments.xml' not in rels_xml:
                # 添加批注关系
                rels_xml = rels_xml.replace(
                    '</Relationships>',
                    '''    <Relationship Id="rIdComments" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>
</Relationships>'''
                )
            
            return rels_xml
            
        except Exception as e:
            print(f"创建关系文件失败: {e}")
            # 返回基本的关系文件
            return '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rIdComments" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>
</Relationships>'''

    def _add_comment_references_to_document(self, document_xml: str, comments_data: list) -> str:
        """在document.xml中添加批注引用标记"""
        try:
            modified_xml = document_xml
            
            # 为每个批注添加引用标记
            for comment in comments_data:
                comment_id = comment['id']
                # 从synchronized_changes中获取原始文本信息
                # 这里需要改进获取原始文本的方法
                
                # 简单的方法：在每个包含跟踪更改的段落后添加批注引用
                import re
                
                # 查找包含删除标记的位置，在其后添加批注引用
                del_pattern = f'(<w:del[^>]*w:id="{comment_id}"[^>]*>.*?</w:del>)'
                matches = re.finditer(del_pattern, modified_xml, re.DOTALL)
                
                for match in matches:
                    # 在删除标记后添加批注标记
                    comment_range_start = f'<w:commentRangeStart w:id="{comment_id}"/>'
                    comment_range_end = f'<w:commentRangeEnd w:id="{comment_id}"/>'
                    comment_reference = f'<w:r><w:commentReference w:id="{comment_id}"/></w:r>'
                    
                    replacement = f'{comment_range_start}{match.group(0)}{comment_range_end}{comment_reference}'
                    modified_xml = modified_xml.replace(match.group(0), replacement, 1)
                    self.console.print(f"[green]✅ 添加批注引用标记: comment_id={comment_id}[/green]")
                    break  # 只处理第一个匹配
            
            return modified_xml
            
        except Exception as e:
            self.console.print(f"[red]添加批注引用失败: {e}[/red]")
            import traceback
            traceback.print_exc()
            return document_xml

    def _prepare_comments_with_changes(self, comments_data: list, synchronized_changes: list) -> list:
        """准备包含更改信息的批注数据"""
        try:
            enhanced_comments = []
            
            for i, comment in enumerate(comments_data):
                enhanced_comment = comment.copy()
                
                # 从synchronized_changes中找到对应的更改信息
                if i < len(synchronized_changes):
                    change = synchronized_changes[i]
                    enhanced_comment['original_text'] = change.get('original_text', '')
                    enhanced_comment['corrected_text'] = change.get('corrected_text', '')
                
                enhanced_comments.append(enhanced_comment)
            
            return enhanced_comments
            
        except Exception as e:
            self.console.print(f"[red]准备批注数据失败: {e}[/red]")
            return comments_data

    def _create_updated_content_types(self, input_zip) -> str:
        """创建更新的内容类型文件"""
        try:
            # 读取原始内容类型文件
            if '[Content_Types].xml' in input_zip.namelist():
                content_types_xml = input_zip.read('[Content_Types].xml').decode('utf-8')
            else:
                content_types_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
</Types>'''
            
            # 检查是否已包含批注内容类型
            if 'word/comments.xml' not in content_types_xml:
                # 添加批注内容类型
                content_types_xml = content_types_xml.replace(
                    '</Types>',
                    '''    <Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>
</Types>'''
                )
            
            return content_types_xml
            
        except Exception as e:
            print(f"创建内容类型文件失败: {e}")
            # 返回基本的内容类型文件
            return '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>
</Types>'''

    def _extract_word_corrections(self, original_text: str, suggested_text: str):
        """从句子级别的修正中提取词汇级别的修正"""
        corrections = []
        try:
            # 简单的词汇差异检测
            original_words = original_text.split()
            suggested_words = suggested_text.split()
            
            # 如果长度相同，逐词比较
            if len(original_words) == len(suggested_words):
                for orig, sugg in zip(original_words, suggested_words):
                    if orig != sugg:
                        corrections.append((orig, sugg))
            else:
                # 如果长度不同，查找明显的替换
                common_replacements = [
                    ("计算器科学", "计算机科学"),
                    ("程式设计", "程序设计"), 
                    ("软体工程", "软件工程"),
                    ("变数", "变量"),
                    ("函式", "函数"),
                    ("超级计算器", "超级计算机"),
                    (",", "，"),  # 标点符号替换
                ]
                
                for orig, repl in common_replacements:
                    if orig in original_text and repl in suggested_text:
                        corrections.append((orig, repl))
            
            return corrections
        except Exception as e:
            self.console.print(f"[yellow]⚠️  提取词汇修正失败: {e}[/yellow]")
            return []

    def _extract_terms_from_inconsistency(self, problem_text: str, suggestion: str):
        """从术语不一致问题中提取术语对"""
        terms = []
        
        # 解析不一致术语描述
        if "发现多种术语：" in problem_text:
            # 提取术语列表
            terms_part = problem_text.split("发现多种术语：")[1].strip()
            # 移除可能的额外描述
            if "，" in terms_part:
                terms_part = terms_part.split("，")[0]
            if "。" in terms_part:
                terms_part = terms_part.split("。")[0]
            
            # 分割术语
            term_variants = []
            if "、" in terms_part:
                term_variants = [t.strip().strip('"').strip("'") for t in terms_part.split("、")]
            elif "，" in terms_part:
                term_variants = [t.strip().strip('"').strip("'") for t in terms_part.split("，")]
            else:
                # 单个术语的情况
                term_variants = [terms_part.strip().strip('"').strip("'")]
            
            # 从建议中提取标准术语
            standard_term = None
            if "建议统一使用" in suggestion:
                standard_part = suggestion.split("建议统一使用")[1].strip()
                if "。" in standard_part:
                    standard_part = standard_part.split("。")[0]
                if "，" in standard_part:
                    standard_part = standard_part.split("，")[0]
                standard_term = standard_part.strip().strip('"').strip("'")
            elif "推荐使用" in suggestion:
                standard_part = suggestion.split("推荐使用")[1].strip()
                if "。" in standard_part:
                    standard_part = standard_part.split("。")[0]
                if "，" in standard_part:
                    standard_part = standard_part.split("，")[0]
                standard_term = standard_part.strip().strip('"').strip("'")
            
            # 如果找到标准术语，为每个变体创建修正对
            if standard_term and term_variants:
                for variant in term_variants:
                    if variant and variant != standard_term:
                        terms.append((variant, standard_term))
                        self.console.print(f"[cyan]📝 术语修正: {variant} → {standard_term}[/cyan]")
            
            # 如果没有明确的标准术语，使用第一个作为标准
            elif len(term_variants) > 1:
                standard_term = term_variants[0]
                for variant in term_variants[1:]:
                    if variant and variant != standard_term:
                        terms.append((variant, standard_term))
                        self.console.print(f"[cyan]📝 术语修正: {variant} → {standard_term}[/cyan]")
        
        # 处理特殊的术语对
        special_corrections = {
            "软体工程": "软件工程",
            "程式设计": "程序设计", 
            "计算器科学": "计算机科学",
            "资料结构": "数据结构",
            "演算法": "算法"
        }
        
        # 检查是否包含特殊术语
        for original, corrected in special_corrections.items():
            if original in problem_text or original in suggestion:
                terms.append((original, corrected))
                self.console.print(f"[cyan]🔧 特殊术语修正: {original} → {corrected}[/cyan]")
        
        return terms

    def _extract_corrected_text(self, suggestion: str):
        """从建议中提取修正后的文本"""
        if not suggestion:
            return None
        
        # 常见的修正模式
        patterns = [
            r"应为[：:]?\s*[\"']([^\"']+)[\"']",
            r"改为[：:]?\s*[\"']([^\"']+)[\"']", 
            r"修正为[：:]?\s*[\"']([^\"']+)[\"']",
            r"建议改为[：:]?\s*[\"']([^\"']+)[\"']",
            r"应该是[：:]?\s*[\"']([^\"']+)[\"']",
            r"正确的是[：:]?\s*[\"']([^\"']+)[\"']",
            r"→\s*[\"']([^\"']+)[\"']",
            r"替换为[：:]?\s*[\"']([^\"']+)[\"']"
        ]
        
        import re
        for pattern in patterns:
            match = re.search(pattern, suggestion)
            if match:
                corrected = match.group(1).strip()
                self.console.print(f"[cyan]🔍 提取修正文本: {corrected}[/cyan]")
                return corrected
        
        # 如果没有找到引号包围的文本，尝试其他模式
        simple_patterns = [
            r"应为[：:]?\s*([^\s，。]+)",
            r"改为[：:]?\s*([^\s，。]+)",
            r"修正为[：:]?\s*([^\s，。]+)",
            r"建议改为[：:]?\s*([^\s，。]+)",
            r"应该是[：:]?\s*([^\s，。]+)",
            r"正确的是[：:]?\s*([^\s，。]+)"
        ]
        
        for pattern in simple_patterns:
            match = re.search(pattern, suggestion)
            if match:
                corrected = match.group(1).strip()
                self.console.print(f"[cyan]🔍 提取修正文本: {corrected}[/cyan]")
                return corrected
        
        # 处理特殊情况：直接的替换建议
        if "→" in suggestion:
            parts = suggestion.split("→")
            if len(parts) >= 2:
                corrected = parts[-1].strip().strip('"').strip("'").strip("。").strip("，")
                if corrected:
                    self.console.print(f"[cyan]🔍 提取修正文本: {corrected}[/cyan]")
                    return corrected
        
        self.console.print(f"[yellow]⚠️  无法提取修正文本: {suggestion}[/yellow]")
        return None
    
    def extract_text_content(self, doc: Document):
        """提取文档的文本内容"""
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        return text_content


def test_fixed_enhanced_proofreader():
    """测试修复版增强校对器"""
    try:
        # 使用测试API密钥
        api_key = "sk-test"
        
        proofreader = ProofReaderWithTrackChangesAndCommentsFixed(api_key)
        
        input_file = "sample_input.docx"
        output_file = "output_fixed.docx"
        
        if os.path.exists(input_file):
            success = proofreader.proofread_with_track_changes_and_comments(input_file, output_file)
            if success:
                print(f"✅ 修复版增强校对成功: {output_file}")
            else:
                print("❌ 修复版增强校对失败")
        else:
            print(f"❌ 输入文件不存在: {input_file}")
            
    except Exception as e:
        print(f"❌ 测试失败: {e}")


if __name__ == "__main__":
    test_fixed_enhanced_proofreader() 
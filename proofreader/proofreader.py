"""
主校对器模块
"""

import os
import sys
from typing import Optional
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.table import Table
from docx import Document
from datetime import datetime

from .config import Config
from .document import DocumentProcessor
from .ai_checker import AIChecker, ProofreadingResult
from .word_comments_advanced import WordCommentsManager
from create_word_comments_xml import add_comments_to_docx

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


class ProofReader:
    """主校对器类"""
    
    def __init__(self, api_key: str = None):
        """初始化校对器"""
        self.config = Config()
        self.ai_checker = AIChecker(api_key)
        self.use_word_review_comments = True  # 启用Word审阅批注
        self.console = Console()
        
        self.document_processor = DocumentProcessor()
    
    def proofread_document(self, input_file: str, output_file: str = None) -> bool:
        """校对文档并生成带批注的输出"""
        try:
            # 生成输出文件名
            if not output_file:
                output_file = input_file.replace('.docx', '_proofread.docx')
            
            self.console.print(f"[green]开始校对文档：{input_file}[/green]")
            
            # 读取文档
            doc = Document(input_file)
            
            # 创建Word批注管理器
            comments_manager = WordCommentsManager(doc)
            
            # 提取文本内容
            text_content = self.extract_text_content(doc)
            self.console.print(f"[blue]提取文本内容: {len(text_content)} 个段落[/blue]")
            
            # 进行AI校对
            self.console.print("[bold]开始AI校对...")
            ai_result = self.ai_checker.check_text(' '.join(text_content))
            
            # 转换AI校对结果为错误列表格式
            errors = self._convert_ai_result_to_errors(ai_result, text_content)
            self.console.print(f"[green]✅ AI校对完成，发现 {len(errors)} 个问题[/green]")
            
            # 添加批注和修正
            comments_data = self.add_comments_and_corrections(doc, errors, comments_manager)
            
            # 完成文档处理
            comments_manager.finalize_document()
            
            # 保存临时文档
            temp_file = output_file.replace('.docx', '_temp.docx')
            doc.save(temp_file)
            
            # 使用完整的Word审阅批注功能
            if add_comments_to_docx(temp_file, output_file, comments_data):
                # 删除临时文件
                os.remove(temp_file)
                self.console.print(f"[green]✅ 校对完成，输出文件: {output_file}[/green]")
                self.console.print("[blue]📝 文档包含完整的Word审阅批注，可在Microsoft Word中查看[/blue]")
                return True
            else:
                # 如果失败，使用临时文件作为输出
                os.rename(temp_file, output_file)
                self.console.print(f"[yellow]⚠️ 审阅批注添加失败，使用基础版本: {output_file}[/yellow]")
                return True
            
        except Exception as e:
            self.console.print(f"[red]❌ 校对失败: {e}[/red]")
            return False
    
    def add_comments_and_corrections(self, doc: Document, errors: list, comments_manager: WordCommentsManager):
        """添加批注和修正，返回批注数据用于完整的Word审阅批注"""
        comment_count = 0
        comments_data = []
        
        for error in errors:
            paragraph_index = error.get('paragraph_index', 0)
            text = error.get('text', '')
            suggestion = error.get('suggestion', '')
            reason = error.get('reason', '')
            
            # 构建批注内容
            comment_text = f"{suggestion}"
            if reason:
                comment_text += f"\n理由: {reason}"
            
            # 获取对应段落
            if paragraph_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[paragraph_index]
                
                # 使用Word审阅批注功能
                if comments_manager.add_comment(paragraph, text, comment_text):
                    comment_count += 1
                    # 添加到批注数据列表
                    comments_data.append({
                        'id': comment_count,
                        'text': comment_text,
                        'author': 'AI校对助手',
                        'date': datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
                    })
                    self.console.print(f"[green]✅ 添加Word审阅批注 {comment_count}: {text} -> {suggestion}[/green]")
                else:
                    self.console.print(f"[red]❌ 批注添加失败: {text}[/red]")
        
        self.console.print(f"[blue]📝 总共添加了 {comment_count} 个Word审阅批注[/blue]")
        return comments_data
    
    def _add_comments_to_document(self, segment: str, result: ProofreadingResult):
        """将校对结果添加为文档批注"""
        for issue in result.issues:
            # 查找包含问题文本的段落
            para_index, para_text = self.document_processor.get_paragraph_by_text(issue["text"])
            
            if para_index >= 0:
                comment = f"{issue['type']}: {issue['suggestion']}"
                self.document_processor.add_comment(
                    para_index, 
                    issue["text"], 
                    comment,
                    self.config.comment_style.author
                )
        
        # 添加改进建议
        for suggestion in result.suggestions:
            para_index, para_text = self.document_processor.get_paragraph_by_text(suggestion["original"])
            
            if para_index >= 0:
                comment = f"建议修改: {suggestion['suggested']} (理由: {suggestion['reason']})"
                self.document_processor.add_comment(
                    para_index,
                    suggestion["original"],
                    comment,
                    self.config.comment_style.author
                )
    
    def _show_report(self, results: list):
        """显示校对报告"""
        self.console.print("\n[bold blue]校对报告[/bold blue]")
        
        # 统计各类问题
        issue_counts = {}
        severity_counts = {"high": 0, "medium": 0, "low": 0}
        
        for segment, result in results:
            for issue in result.issues:
                issue_type = issue["type"]
                severity = issue["severity"]
                
                issue_counts[issue_type] = issue_counts.get(issue_type, 0) + 1
                severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        # 创建问题类型统计表
        if issue_counts:
            table = Table(title="问题类型统计")
            table.add_column("问题类型", style="cyan")
            table.add_column("数量", style="magenta")
            
            for issue_type, count in sorted(issue_counts.items(), key=lambda x: x[1], reverse=True):
                table.add_row(issue_type, str(count))
            
            self.console.print(table)
        
        # 创建严重程度统计表
        severity_table = Table(title="问题严重程度统计")
        severity_table.add_column("严重程度", style="cyan")
        severity_table.add_column("数量", style="magenta")
        severity_table.add_column("颜色标识", style="white")
        
        colors = {"high": "[red]高[/red]", "medium": "[yellow]中[/yellow]", "low": "[green]低[/green]"}
        for severity, count in severity_counts.items():
            if count > 0:
                severity_table.add_row(severity, str(count), colors[severity])
        
        self.console.print(severity_table)
        
        # 显示详细问题列表
        self._show_detailed_issues(results)
    
    def _show_detailed_issues(self, results: list):
        """显示详细问题列表"""
        self.console.print("\n[bold blue]详细问题列表[/bold blue]")
        
        issue_num = 1
        for segment, result in results:
            if result.issues:
                self.console.print(f"\n[bold cyan]段落内容:[/bold cyan] {segment[:100]}...")
                
                for issue in result.issues:
                    severity_color = {"high": "red", "medium": "yellow", "low": "green"}
                    color = severity_color.get(issue["severity"], "white")
                    
                    self.console.print(f"[{color}]{issue_num}. {issue['type']}[/{color}]")
                    self.console.print(f"   问题文本: {issue['text']}")
                    self.console.print(f"   修改建议: {issue['suggestion']}")
                    self.console.print(f"   严重程度: {issue['severity']}")
                    
                    issue_num += 1
                    
                    if issue_num > 20:  # 限制显示数量
                        self.console.print(f"[yellow]... 还有更多问题，请查看输出文档中的批注[/yellow]")
                        return
    
    def batch_proofread(self, input_dir: str, output_dir: str) -> bool:
        """批量校对目录下的所有docx文件"""
        try:
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            docx_files = [f for f in os.listdir(input_dir) if f.endswith('.docx')]
            
            if not docx_files:
                self.console.print("[yellow]未找到docx文件[/yellow]")
                return False
            
            self.console.print(f"[green]找到 {len(docx_files)} 个文档待校对[/green]")
            
            success_count = 0
            for filename in docx_files:
                input_path = os.path.join(input_dir, filename)
                output_path = os.path.join(output_dir, f"校对_{filename}")
                
                self.console.print(f"\n[blue]处理文件：{filename}[/blue]")
                
                if self.proofread_document(input_path, output_path):
                    success_count += 1
                else:
                    self.console.print(f"[red]文件 {filename} 校对失败[/red]")
            
            self.console.print(f"\n[green]批量校对完成！成功处理 {success_count}/{len(docx_files)} 个文件[/green]")
            return success_count == len(docx_files)
            
        except Exception as e:
            self.console.print(f"[red]批量校对失败：{e}[/red]")
            return False
    
    def quick_check(self, text: str) -> ProofreadingResult:
        """快速检查文本片段"""
        return self.ai_checker.check_text(text)

    def extract_text_content(self, doc: Document):
        """提取文档的文本内容"""
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        return text_content

    def _convert_ai_result_to_errors(self, ai_result: ProofreadingResult, text_content: list):
        """将AI校对结果转换为错误列表格式"""
        errors = []
        
        # 处理issues
        for issue in ai_result.issues:
            # 找到问题文本在哪个段落
            problem_text = issue.get('text', '')
            for i, paragraph_text in enumerate(text_content):
                if problem_text in paragraph_text:
                    errors.append({
                        'paragraph_index': i,
                        'text': problem_text,
                        'suggestion': issue.get('suggestion', ''),
                        'reason': f"{issue.get('type', '')} - {issue.get('severity', '')}"
                    })
                    break
        
        # 处理suggestions
        for suggestion in ai_result.suggestions:
            original_text = suggestion.get('original', '')
            for i, paragraph_text in enumerate(text_content):
                if original_text in paragraph_text:
                    errors.append({
                        'paragraph_index': i,
                        'text': original_text,
                        'suggestion': suggestion.get('suggested', ''),
                        'reason': suggestion.get('reason', '')
                    })
                    break
        
        return errors 
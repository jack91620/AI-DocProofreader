#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
增强版校对器 - 同时使用Word跟踪更改和批注功能
"""

import os
import sys
from typing import Optional
from rich.console import Console
from docx import Document
from datetime import datetime

from .config import Config
from .document import DocumentProcessor
from .ai_checker import AIChecker, ProofreadingResult
from .word_track_changes import WordTrackChangesManager, enable_track_changes_in_docx
from .word_comments_advanced import WordCommentsManager
from .word_comments_xml import create_comments_xml, create_document_rels, update_content_types, add_comments_to_docx
import zipfile
import tempfile


class ProofReaderWithTrackChangesAndComments:
    """增强版校对器 - 同时使用跟踪更改和批注"""
    
    def __init__(self, api_key: str = None):
        """初始化校对器"""
        self.config = Config()
        if api_key:
            self.config.ai.api_key = api_key
        self.ai_checker = AIChecker(self.config)
        self.console = Console()
        
        self.document_processor = DocumentProcessor()
    
    def proofread_with_track_changes_and_comments(self, input_file: str, output_file: str = None) -> bool:
        """使用跟踪更改和批注进行校对 - 确保跟踪更改和批注完全同步"""
        try:
            # 生成输出文件名
            if not output_file:
                output_file = input_file.replace('.docx', '_tracked_with_comments.docx')
            
            self.console.print(f"[green]开始增强校对：{input_file}[/green]")
            
            # 第一步：进行AI校对获取所有问题
            self.console.print("[blue]第一步：AI校对分析文档...[/blue]")
            doc = Document(input_file)
            text_content = self.extract_text_content(doc)
            self.console.print(f"[blue]提取文本内容: {len(text_content)} 个段落[/blue]")
            
            # 只进行一次AI校对
            self.console.print("[bold]开始AI校对...")
            ai_result = self.ai_checker.check_text(' '.join(text_content))
            
            # 第二步：同时创建跟踪更改和批注的数据
            self.console.print("[blue]第二步：同步生成跟踪更改和批注数据...[/blue]")
            synchronized_changes = self._create_synchronized_changes(ai_result, text_content)
            self.console.print(f"[green]✅ AI校对完成，发现 {len(synchronized_changes)} 个问题[/green]")
            
            # 第三步：同时应用跟踪更改和批注
            self.console.print("[blue]第三步：同时应用跟踪更改和批注...[/blue]")
            success = self._apply_synchronized_changes(doc, synchronized_changes, output_file)
            
            if success:
                self.console.print(f"[green]✅ 增强校对完成：{output_file}[/green]")
                self.console.print("[blue]📝 文档包含：[/blue]")
                self.console.print("   - 🔄 真正的Word跟踪更改")
                self.console.print("   - 💬 每个修订都有对应的详细批注")
                self.console.print("   - ✅ 完全同步的修订和批注系统")
                return True
            else:
                return False
            
        except Exception as e:
            self.console.print(f"[red]❌ 增强校对失败: {e}[/red]")
            return False

    def _create_synchronized_changes(self, ai_result: ProofreadingResult, text_content: list):
        """创建同步的跟踪更改和批注数据"""
        synchronized_changes = []
        
        # 处理AI校对的issues
        for issue in ai_result.issues:
            problem_text = issue.get('text', '')
            suggestion = issue.get('suggestion', '')
            issue_type = issue.get('type', '')
            severity = issue.get('severity', '')
            
            # 提取修正后的文本
            corrected_text = self._extract_corrected_text(suggestion)
            
            # 如果有可用的修正文本，创建同步更改
            if corrected_text and corrected_text != problem_text:
                # 找到问题文本在哪个段落
                for i, paragraph_text in enumerate(text_content):
                    if problem_text in paragraph_text:
                        # 创建批注文本
                        comment_text = f"🔍 发现问题: {issue_type}\n"
                        comment_text += f"📝 修正: {problem_text} → {corrected_text}\n"
                        comment_text += f"⚠️ 严重程度: {severity}\n"
                        comment_text += f"💡 建议: {suggestion}\n"
                        comment_text += f"⏰ 检查时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                        
                        synchronized_changes.append({
                            'paragraph_index': i,
                            'original_text': problem_text,
                            'corrected_text': corrected_text,
                            'comment_text': comment_text,
                            'reason': f"{issue_type} - {severity}",
                            'type': 'issue_fix'
                        })
                        break
        
        # 处理AI校对的suggestions
        for suggestion in ai_result.suggestions:
            original_text = suggestion.get('original', '')
            suggested_text = suggestion.get('suggested', '')
            reason = suggestion.get('reason', '')
            
            # 如果建议文本与原文本不同，创建同步更改
            if suggested_text and suggested_text != original_text:
                # 找到原文本在哪个段落
                for i, paragraph_text in enumerate(text_content):
                    if original_text in paragraph_text:
                        # 创建批注文本
                        comment_text = f"💡 建议修改: '{original_text}' → '{suggested_text}'\n"
                        comment_text += f"📋 原因: {reason}\n"
                        comment_text += f"🎯 类型: 改进建议\n"
                        comment_text += f"⏰ 建议时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                        
                        synchronized_changes.append({
                            'paragraph_index': i,
                            'original_text': original_text,
                            'corrected_text': suggested_text,
                            'comment_text': comment_text,
                            'reason': reason,
                            'type': 'suggestion'
                        })
                        break
        
        return synchronized_changes

    def _apply_synchronized_changes(self, doc: Document, synchronized_changes: list, output_file: str) -> bool:
        """同时应用跟踪更改和批注"""
        try:
            # 创建跟踪更改管理器
            track_changes_manager = WordTrackChangesManager(doc)
            
            # 创建批注管理器
            comments_manager = WordCommentsManager(doc)
            
            # 应用每个同步更改
            applied_count = 0
            for change in synchronized_changes:
                paragraph_index = change.get('paragraph_index', 0)
                original_text = change.get('original_text', '')
                corrected_text = change.get('corrected_text', '')
                comment_text = change.get('comment_text', '')
                reason = change.get('reason', '')
                change_type = change.get('type', '')
                
                # 获取对应段落
                if paragraph_index < len(doc.paragraphs):
                    paragraph = doc.paragraphs[paragraph_index]
                    
                    # 同时应用跟踪更改和批注
                    track_change_success = track_changes_manager.add_tracked_change(
                        paragraph, original_text, corrected_text, reason
                    )
                    comment_success = comments_manager.add_comment(
                        paragraph, original_text, comment_text
                    )
                    
                    if track_change_success and comment_success:
                        applied_count += 1
                        self.console.print(
                            f"[green]✅ 同步更改 {applied_count}: {original_text} -> {corrected_text} + 批注[/green]"
                        )
                    elif track_change_success:
                        self.console.print(
                            f"[yellow]⚠️ 跟踪更改成功但批注失败: {original_text}[/yellow]"
                        )
                    elif comment_success:
                        self.console.print(
                            f"[yellow]⚠️ 批注成功但跟踪更改失败: {original_text}[/yellow]"
                        )
                    else:
                        self.console.print(
                            f"[red]❌ 同步更改失败: {original_text}[/red]"
                        )
            
            # 应用所有跟踪更改
            track_changes_manager.apply_all_changes()
            
            # 保存临时文档
            temp_file = output_file.replace('.docx', '_temp.docx')
            doc.save(temp_file)
            
            # 生成最终文档（包含跟踪更改和批注）
            success = self._create_final_synchronized_document(
                temp_file, 
                output_file, 
                track_changes_manager.revisions_data,
                comments_manager.get_comments_for_xml()
            )
            
            # 清理临时文件
            if os.path.exists(temp_file):
                os.remove(temp_file)
            
            if success:
                self.console.print(f"[green]✅ 成功应用 {applied_count} 个同步更改（跟踪更改+批注）[/green]")
                return True
            else:
                return False
            
        except Exception as e:
            self.console.print(f"[red]❌ 应用同步更改失败: {e}[/red]")
            return False

    def _create_final_synchronized_document(self, temp_file: str, output_file: str, track_changes_data: list, comments_data: list) -> bool:
        """创建最终的同步文档（包含跟踪更改和批注）"""
        try:
            # 导入所需模块
            from .word_track_changes import enable_track_changes_in_docx
            from .word_track_changes_with_comments import ProofReaderWithCommentsAndTrackChanges
            
            # 使用word_track_changes_with_comments模块来处理批注和跟踪更改
            # 这个模块能够正确处理批注引用标记
            comment_proofreader = ProofReaderWithCommentsAndTrackChanges()
            
            # 格式化批注数据为正确的格式
            formatted_comments = []
            for comment in comments_data:
                formatted_comments.append({
                    'text': comment.get('text', ''),
                    'author': comment.get('author', 'AI校对助手'),
                    'date': comment.get('date', datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"))
                })
            
            self.console.print(f"[cyan]🔧 使用专门的批注处理模块添加 {len(formatted_comments)} 个批注[/cyan]")
            
            # 直接使用专门的批注处理方法
            success = comment_proofreader._add_comments_to_docx_with_references(
                temp_file, 
                output_file, 
                formatted_comments
            )
            
            if success:
                self.console.print("[green]✅ 成功创建包含跟踪更改和批注的文档[/green]")
                return True
            else:
                self.console.print("[red]❌ 创建最终文档失败[/red]")
                return False
                
        except Exception as e:
            self.console.print(f"[red]❌ 创建最终同步文档失败: {e}[/red]")
            
            # 回退方案：仅使用跟踪更改
            try:
                self.console.print("[yellow]⚠️ 尝试回退方案：仅包含跟踪更改[/yellow]")
                from .word_track_changes import enable_track_changes_in_docx
                
                success = enable_track_changes_in_docx(temp_file, output_file, track_changes_data)
                if success:
                    self.console.print("[yellow]✅ 回退方案成功：文档包含跟踪更改，但缺少批注[/yellow]")
                    return True
                else:
                    self.console.print("[red]❌ 回退方案也失败了[/red]")
                    return False
                    
            except Exception as fallback_error:
                self.console.print(f"[red]❌ 回退方案失败: {fallback_error}[/red]")
                return False

    def _extract_corrected_text(self, suggestion: str):
        """从建议中提取修正后的文本"""
        # 尝试从建议中提取修正文本
        if "建议改为：" in suggestion:
            return suggestion.split("建议改为：")[-1].strip()
        elif "应为" in suggestion:
            return suggestion.split("应为")[-1].strip().strip("'\"")
        elif "->" in suggestion:
            return suggestion.split("->")[-1].strip()
        elif "改为" in suggestion:
            return suggestion.split("改为")[-1].strip().strip("'\"")
        else:
            # 如果无法提取，返回空字符串
            return ""
    
    def extract_text_content(self, doc: Document):
        """提取文档的文本内容"""
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        return text_content


# 测试函数
def test_enhanced_proofreader():
    """测试增强版校对器"""
    try:
        # 使用测试API密钥
        api_key = "sk-test"  # 替换为真实的API密钥
        
        proofreader = ProofReaderWithTrackChangesAndComments(api_key)
        
        input_file = "sample_input.docx"
        output_file = "sample_output_enhanced_track_changes_comments.docx"
        
        if os.path.exists(input_file):
            success = proofreader.proofread_with_track_changes_and_comments(input_file, output_file)
            if success:
                print(f"✅ 增强校对成功: {output_file}")
            else:
                print("❌ 增强校对失败")
        else:
            print(f"❌ 输入文件不存在: {input_file}")
            
    except Exception as e:
        print(f"❌ 测试失败: {e}")


if __name__ == "__main__":
    test_enhanced_proofreader() 
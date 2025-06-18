#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
带批注的真正Word跟踪更改功能模块
实现Microsoft Word审阅中的修订功能，同时添加批注说明修订原因
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import zipfile
import tempfile
import os
import xml.etree.ElementTree as ET
try:
    from .word_track_changes import WordTrackChangesManager, enable_track_changes_in_docx
    from .word_comments_advanced import WordCommentsManager
except ImportError:
    from word_track_changes import WordTrackChangesManager, enable_track_changes_in_docx
    from word_comments_advanced import WordCommentsManager


class WordTrackChangesWithCommentsManager:
    """带批注的Word跟踪更改管理器 - 同时生成修订和批注"""
    
    def __init__(self, document):
        self.document = document
        self.track_changes_manager = WordTrackChangesManager(document)
        self.comments_manager = WordCommentsManager(document)
        self.combined_changes = []
    
    def add_tracked_change_with_comment(self, paragraph, original_text: str, corrected_text: str, reason: str = ""):
        """添加跟踪更改并同时添加批注说明原因"""
        try:
            # 1. 先为原文添加批注（避免修订后文本查找问题）
            comment_text = self._generate_comment_text(original_text, corrected_text, reason)
            comment_success = self.comments_manager.add_comment(
                paragraph, original_text, comment_text
            )
            
            # 2. 再添加跟踪更改
            track_success = self.track_changes_manager.add_tracked_change(
                paragraph, original_text, corrected_text, reason
            )
            
            if not track_success:
                print(f"⚠️ 跟踪更改添加失败: {original_text}")
                return comment_success  # 即使跟踪更改失败，如果批注成功也算部分成功
            
            # 3. 记录组合修改
            self.combined_changes.append({
                'original_text': original_text,
                'corrected_text': corrected_text,
                'reason': reason,
                'track_changes_success': track_success,
                'comment_success': comment_success,
                'paragraph': paragraph
            })
            
            print(f"✅ 已添加跟踪更改+批注: {original_text} -> {corrected_text}")
            if comment_success:
                print(f"   📝 批注内容: {comment_text}")
            else:
                print(f"   ⚠️  批注添加失败，仅应用跟踪更改")
            
            return True
            
        except Exception as e:
            print(f"添加跟踪更改+批注失败: {e}")
            return False
    
    def _generate_comment_text(self, original_text: str, corrected_text: str, reason: str):
        """生成批注文本"""
        comment_parts = []
        
        # 修订说明
        comment_parts.append(f"🔄 修订: '{original_text}' → '{corrected_text}'")
        
        # 修订原因
        if reason:
            comment_parts.append(f"📋 原因: {reason}")
        
        # 修订类型判断
        revision_type = self._determine_revision_type(original_text, corrected_text)
        comment_parts.append(f"🏷️ 类型: {revision_type}")
        
        # 时间戳
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        comment_parts.append(f"⏰ 时间: {timestamp}")
        
        return "\n".join(comment_parts)
    
    def _determine_revision_type(self, original_text: str, corrected_text: str):
        """判断修订类型"""
        if len(original_text) == 1 and len(corrected_text) == 1:
            return "错别字修正"
        elif "科学" in original_text or "科学" in corrected_text:
            return "术语修正"
        elif len(original_text) > len(corrected_text):
            return "文本简化"
        elif len(original_text) < len(corrected_text):
            return "文本扩展"
        else:
            return "文本优化"
    
    def apply_all_changes(self):
        """应用所有跟踪更改"""
        print("🔄 应用所有跟踪更改...")
        self.track_changes_manager.apply_all_changes()
        print("✅ 跟踪更改应用完成")
    
    def get_statistics(self):
        """获取修订统计信息"""
        total_changes = len(self.combined_changes)
        track_changes_count = sum(1 for c in self.combined_changes if c['track_changes_success'])
        comments_count = sum(1 for c in self.combined_changes if c['comment_success'])
        
        return {
            'total_changes': total_changes,
            'track_changes_count': track_changes_count,
            'comments_count': comments_count,
            'success_rate': (track_changes_count / total_changes * 100) if total_changes > 0 else 0
        }


def enable_track_changes_and_comments_in_docx(docx_path, output_path, track_changes_data, comments_data):
    """在Word文档中启用跟踪更改和批注"""
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            # 解压docx文件
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # 启用跟踪更改
            enable_track_changes_setting(temp_dir)
            
            # 添加批注XML文件
            add_comments_xml_files(temp_dir, comments_data)
            
            # 验证XML结构
            verify_combined_xml_structure(temp_dir)
            
            # 重新打包
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arc_name = os.path.relpath(file_path, temp_dir)
                        zip_ref.write(file_path, arc_name)
            
            print(f"✅ 成功创建带批注的Word跟踪更改文档: {output_path}")
            return True
            
    except Exception as e:
        print(f"❌ 创建带批注的跟踪更改文档失败: {e}")
        return False


def enable_track_changes_setting(temp_dir):
    """启用跟踪更改设置"""
    try:
        settings_path = os.path.join(temp_dir, 'word', 'settings.xml')
        
        if os.path.exists(settings_path):
            tree = ET.parse(settings_path)
            root = tree.getroot()
        else:
            ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            root = ET.Element(f'{{{ns_w}}}settings')
            tree = ET.ElementTree(root)
        
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # 启用跟踪更改
        track_revisions = root.find('.//w:trackRevisions', ns)
        if track_revisions is None:
            track_revisions = ET.SubElement(root, f"{{{ns['w']}}}trackRevisions")
        track_revisions.set(f"{{{ns['w']}}}val", "1")
        
        # 确保目录存在
        word_dir = os.path.join(temp_dir, 'word')
        os.makedirs(word_dir, exist_ok=True)
        
        tree.write(settings_path, encoding='utf-8', xml_declaration=True)
        print("✅ 已启用跟踪更改设置")
        
    except Exception as e:
        print(f"❌ 启用跟踪更改设置失败: {e}")


def add_comments_xml_files(temp_dir, comments_data):
    """添加批注相关的XML文件"""
    try:
        try:
            from .word_comments_xml import (
                create_comments_xml, 
                create_document_rels, 
                update_content_types
            )
        except ImportError:
            from .word_comments_xml import (
                create_comments_xml, 
                create_document_rels, 
                update_content_types
            )
        
        word_dir = os.path.join(temp_dir, 'word')
        os.makedirs(word_dir, exist_ok=True)
        
        # 创建comments.xml
        comments_xml_path = os.path.join(word_dir, 'comments.xml')
        create_comments_xml(comments_xml_path, comments_data)
        
        # 创建document.xml.rels
        create_document_rels(temp_dir)
        
        # 更新Content_Types.xml
        update_content_types(temp_dir)
        
        print("✅ 已添加批注XML文件")
        
    except Exception as e:
        print(f"❌ 添加批注XML文件失败: {e}")


def verify_combined_xml_structure(temp_dir):
    """验证组合XML结构"""
    try:
        document_path = os.path.join(temp_dir, 'word', 'document.xml')
        comments_path = os.path.join(temp_dir, 'word', 'comments.xml')
        
        # 验证document.xml中的修订标记
        if os.path.exists(document_path):
            with open(document_path, 'r', encoding='utf-8') as f:
                doc_content = f.read()
            
            del_count = doc_content.count('<w:del ')
            ins_count = doc_content.count('<w:ins ')
            comment_ref_count = doc_content.count('<w:commentReference ')
            
            print(f"📊 document.xml统计:")
            print(f"   - 删除标记: {del_count}")
            print(f"   - 插入标记: {ins_count}")
            print(f"   - 批注引用: {comment_ref_count}")
        
        # 验证comments.xml
        if os.path.exists(comments_path):
            with open(comments_path, 'r', encoding='utf-8') as f:
                comments_content = f.read()
            
            comment_count = comments_content.count('<w:comment ')
            print(f"   - 批注数量: {comment_count}")
            
            if del_count > 0 and comment_count > 0:
                print("✅ 发现跟踪更改和批注标记")
            else:
                print("⚠️  跟踪更改或批注标记缺失")
        
    except Exception as e:
        print(f"❌ 验证XML结构失败: {e}")


# 测试函数
def test_track_changes_with_comments():
    """测试带批注的跟踪更改功能"""
    try:
        print("🔄 开始测试带批注的跟踪更改功能...")
        
        # 创建测试文档
        doc = Document()
        doc.add_heading('测试文档 - 带批注的跟踪更改', 0)
        doc.add_paragraph("这是一个测试段落。")
        doc.add_paragraph("计算器科学是一门非常重要的学科，涉及到程式设计和筭法等内容。")
        doc.add_paragraph("在日常生活中，我们经常需要进行文字校对工作。")
        
        # 创建带批注的跟踪更改管理器
        track_comments_manager = WordTrackChangesWithCommentsManager(doc)
        
        # 添加带批注的跟踪更改
        paragraphs = list(doc.paragraphs)
        
        changes = [
            (paragraphs[2], "计算器科学", "计算机科学", "错别字修正：'器'应为'机'"),
            (paragraphs[2], "程式设计", "程序设计", "术语统一：使用标准中文术语"),
            (paragraphs[2], "筭法", "算法", "错别字修正：'筭'应为'算'"),
        ]
        
        for paragraph, original, corrected, reason in changes:
            track_comments_manager.add_tracked_change_with_comment(
                paragraph, original, corrected, reason
            )
        
        # 应用所有更改
        track_comments_manager.apply_all_changes()
        
        # 获取统计信息
        stats = track_comments_manager.get_statistics()
        print(f"\n📊 修订统计:")
        print(f"   - 总修改数: {stats['total_changes']}")
        print(f"   - 跟踪更改数: {stats['track_changes_count']}")
        print(f"   - 批注数: {stats['comments_count']}")
        print(f"   - 成功率: {stats['success_rate']:.1f}%")
        
        # 保存临时文档
        temp_file = "test_track_changes_with_comments_temp.docx"
        doc.save(temp_file)
        
        # 生成最终文档
        output_file = "test_track_changes_with_comments.docx"
        success = enable_track_changes_and_comments_in_docx(
            temp_file, 
            output_file, 
            track_comments_manager.track_changes_manager.revisions_data,
            track_comments_manager.comments_manager.comments
        )
        
        if success:
            print(f"\n✅ 带批注的跟踪更改文档已创建: {output_file}")
            print("📝 现在可以在Microsoft Word中查看:")
            print("   - 真正的跟踪更改（红色删除线 + 蓝色下划线）")
            print("   - 详细的批注说明（修订原因和类型）")
            print("   - 可以接受/拒绝修改和回复批注")
            
            # 清理临时文件
            if os.path.exists(temp_file):
                os.remove(temp_file)
        else:
            print("❌ 创建失败")
        
        return success
        
    except Exception as e:
        print(f"❌ 测试失败: {e}")
        return False


if __name__ == "__main__":
    test_track_changes_with_comments()


class ProofReaderWithCommentsAndTrackChanges:
    """带批注的校对器 - 同时处理跟踪更改和批注，确保批注引用标记正确保存"""
    
    def __init__(self):
        self.console = None
        try:
            from rich.console import Console
            self.console = Console()
        except ImportError:
            pass
    
    def _print(self, message, style=""):
        """统一的打印方法"""
        if self.console:
            if style:
                self.console.print(f"[{style}]{message}[/{style}]")
            else:
                self.console.print(message)
        else:
            print(message)
    
    def _add_comments_to_docx_with_references(self, input_file: str, output_file: str, comments_data: list) -> bool:
        """直接在document.xml中添加批注引用标记，然后添加comments.xml文件"""
        try:
            self._print(f"🔧 开始处理批注引用标记和XML: {len(comments_data)} 个批注", "cyan")
            
            with tempfile.TemporaryDirectory() as temp_dir:
                # 1. 解压文档
                with zipfile.ZipFile(input_file, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                # 2. 在document.xml中添加批注引用标记
                document_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
                if not self._add_comment_references_to_document_xml(document_xml_path, comments_data):
                    self._print("❌ 添加批注引用标记失败", "red")
                    return False
                
                # 3. 创建comments.xml文件
                comments_xml_path = os.path.join(temp_dir, 'word', 'comments.xml')
                if not self._create_comments_xml_file(comments_xml_path, comments_data):
                    self._print("❌ 创建comments.xml失败", "red")
                    return False
                
                # 4. 更新文档关系和内容类型
                self._update_document_relationships(temp_dir)
                self._update_content_types(temp_dir)
                
                # 5. 重新打包文档
                with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arc_name = os.path.relpath(file_path, temp_dir)
                            zip_ref.write(file_path, arc_name)
                
                self._print(f"✅ 成功创建包含批注的文档: {output_file}", "green")
                return True
                
        except Exception as e:
            self._print(f"❌ 处理批注失败: {e}", "red")
            import traceback
            traceback.print_exc()
            return False
    
    def _add_comment_references_to_document_xml(self, document_xml_path: str, comments_data: list) -> bool:
        """在document.xml中添加批注引用标记"""
        try:
            # 读取document.xml
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            # 定义命名空间
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            ET.register_namespace('w', ns['w'])
            
            # 查找所有段落
            paragraphs = root.findall('.//w:p', ns)
            
            comment_id = 1
            for comment_data in comments_data:
                comment_text_content = comment_data.get('text', '')
                
                # 尝试从批注文本中提取原始文本（用于定位）
                original_text = self._extract_original_text_from_comment(comment_text_content)
                
                if not original_text:
                    self._print(f"⚠️ 无法从批注中提取原始文本: {comment_text_content[:50]}...", "yellow")
                    continue
                
                # 在段落中查找原始文本并添加批注引用
                if self._add_comment_ref_to_paragraph(paragraphs, original_text, comment_id, ns):
                    self._print(f"✅ 已添加批注引用 {comment_id}: {original_text}", "green")
                    comment_id += 1
                else:
                    self._print(f"⚠️ 未找到文本位置: {original_text}", "yellow")
            
            # 保存修改后的document.xml
            tree.write(document_xml_path, encoding='utf-8', xml_declaration=True)
            return True
            
        except Exception as e:
            self._print(f"❌ 修改document.xml失败: {e}", "red")
            return False
    
    def _extract_original_text_from_comment(self, comment_text: str) -> str:
        """从批注文本中提取原始文本"""
        try:
            # 查找 "修订: 'xxx' →" 模式
            if "修订: '" in comment_text and "' →" in comment_text:
                start = comment_text.find("修订: '") + 4
                end = comment_text.find("' →", start)
                if start > 3 and end > start:
                    return comment_text[start:end]
            
            # 查找其他可能的模式
            if "原文：" in comment_text:
                lines = comment_text.split('\n')
                for line in lines:
                    if "原文：" in line:
                        return line.split("原文：")[1].strip()
            
            return ""
        except Exception:
            return ""
    
    def _add_comment_ref_to_paragraph(self, paragraphs, target_text: str, comment_id: int, ns: dict) -> bool:
        """在段落中添加批注引用标记"""
        try:
            for paragraph in paragraphs:
                # 获取段落的文本内容
                para_text = self._get_paragraph_text(paragraph, ns)
                
                if target_text in para_text:
                    # 找到包含目标文本的段落，添加批注引用标记
                    self._insert_comment_markers(paragraph, target_text, comment_id, ns)
                    return True
            
            return False
        except Exception as e:
            self._print(f"添加批注引用标记失败: {e}", "red")
            return False
    
    def _get_paragraph_text(self, paragraph, ns: dict) -> str:
        """获取段落的纯文本内容"""
        text_parts = []
        text_elements = paragraph.findall('.//w:t', ns)
        for text_elem in text_elements:
            if text_elem.text:
                text_parts.append(text_elem.text)
        return ''.join(text_parts)
    
    def _insert_comment_markers(self, paragraph, target_text: str, comment_id: int, ns: dict):
        """在段落中插入批注标记"""
        try:
            # 创建批注范围开始标记
            comment_range_start = ET.Element(f"{{{ns['w']}}}commentRangeStart")
            comment_range_start.set(f"{{{ns['w']}}}id", str(comment_id))
            
            # 创建批注范围结束标记
            comment_range_end = ET.Element(f"{{{ns['w']}}}commentRangeEnd")
            comment_range_end.set(f"{{{ns['w']}}}id", str(comment_id))
            
            # 创建一个新的run包含批注引用
            comment_run = ET.Element(f"{{{ns['w']}}}r")
            comment_ref = ET.SubElement(comment_run, f"{{{ns['w']}}}commentReference")
            comment_ref.set(f"{{{ns['w']}}}id", str(comment_id))
            
            # 将批注标记插入到段落的适当位置
            # 这里简化处理，在段落末尾添加标记
            paragraph.append(comment_range_start)
            paragraph.append(comment_range_end)
            paragraph.append(comment_run)
            
            self._print(f"✅ 批注标记已插入段落: comment_id={comment_id}", "dim")
            
        except Exception as e:
            self._print(f"插入批注标记失败: {e}", "red")
    
    def _create_comments_xml_file(self, comments_xml_path: str, comments_data: list) -> bool:
        """创建comments.xml文件"""
        try:
            try:
                from .word_comments_xml import create_comments_xml
            except ImportError:
                from word_comments_xml import create_comments_xml
            
            # 为comments_data添加ID
            processed_comments = []
            for i, comment in enumerate(comments_data, 1):
                processed_comment = comment.copy()
                processed_comment['id'] = i
                processed_comments.append(processed_comment)
            
            create_comments_xml(comments_xml_path, processed_comments)
            return True
            
        except Exception as e:
            self._print(f"创建comments.xml失败: {e}", "red")
            return False
    
    def _update_document_relationships(self, temp_dir: str):
        """更新文档关系"""
        try:
            try:
                from .word_comments_xml import create_document_rels
            except ImportError:
                from word_comments_xml import create_document_rels
            create_document_rels(temp_dir)
        except Exception as e:
            self._print(f"更新文档关系失败: {e}", "yellow")
    
    def _update_content_types(self, temp_dir: str):
        """更新内容类型"""
        try:
            try:
                from .word_comments_xml import update_content_types
            except ImportError:
                from word_comments_xml import update_content_types
            update_content_types(temp_dir)
        except Exception as e:
            self._print(f"更新内容类型失败: {e}", "yellow") 
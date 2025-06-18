#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
集成修订功能的校对器 - 使用Word跟踪更改显示AI修改
"""

import os
import sys
from typing import Optional
from rich.console import Console
from docx import Document
from datetime import datetime

from .config import Config
from .document import DocumentProcessor
from .ai_checker import AIChecker, ProofreadingResult
from .word_revisions import SimpleWordRevisionsManager


class ProofReaderWithRevisions:
    """带修订功能的校对器"""
    
    def __init__(self, api_key: str = None):
        """初始化校对器"""
        self.config = Config()
        if api_key:
            self.config.ai.api_key = api_key
        self.ai_checker = AIChecker(self.config)
        self.console = Console()
        
        self.document_processor = DocumentProcessor()
    
    def proofread_document_with_revisions(self, input_file: str, output_file: str = None) -> bool:
        """校对文档并使用修订功能显示更改"""
        try:
            # 生成输出文件名
            if not output_file:
                output_file = input_file.replace('.docx', '_revised.docx')
            
            self.console.print(f"[green]开始修订校对：{input_file}[/green]")
            
            # 读取文档
            doc = Document(input_file)
            
            # 创建修订管理器
            revisions_manager = SimpleWordRevisionsManager(doc)
            
            # 提取文本内容
            text_content = self.extract_text_content(doc)
            self.console.print(f"[blue]提取文本内容: {len(text_content)} 个段落[/blue]")
            
            # 进行AI校对
            self.console.print("[bold]开始AI校对...")
            ai_result = self.ai_checker.check_text(' '.join(text_content))
            
            # 转换AI校对结果为修订格式
            revisions = self._convert_ai_result_to_revisions(ai_result, text_content)
            self.console.print(f"[green]✅ AI校对完成，发现 {len(revisions)} 个需要修订的问题[/green]")
            
            # 应用修订
            revision_count = self.apply_revisions(doc, revisions, revisions_manager)
            
            # 保存文档
            doc.save(output_file)
            self.console.print(f"[green]✅ 修订校对完成，输出文件: {output_file}[/green]")
            self.console.print(f"[blue]📝 已应用 {revision_count} 个修订，修改内容在Word中显示为跟踪更改[/blue]")
            
            return True
            
        except Exception as e:
            self.console.print(f"[red]❌ 修订校对失败: {e}[/red]")
            return False
    
    def apply_revisions(self, doc: Document, revisions: list, revisions_manager: SimpleWordRevisionsManager):
        """应用修订到文档"""
        revision_count = 0
        
        for revision in revisions:
            paragraph_index = revision.get('paragraph_index', 0)
            original_text = revision.get('original_text', '')
            corrected_text = revision.get('corrected_text', '')
            reason = revision.get('reason', '')
            
            # 获取对应段落
            if paragraph_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[paragraph_index]
                
                # 应用修订
                if revisions_manager.add_revision(paragraph, original_text, corrected_text, reason):
                    revision_count += 1
                    self.console.print(f"[green]✅ 修订 {revision_count}: {original_text} -> {corrected_text}[/green]")
                else:
                    self.console.print(f"[red]❌ 修订失败: {original_text}[/red]")
        
        return revision_count
    
    def _convert_ai_result_to_revisions(self, ai_result: ProofreadingResult, text_content: list):
        """将AI校对结果转换为修订格式"""
        revisions = []
        
        # 处理issues
        for issue in ai_result.issues:
            problem_text = issue.get('text', '')
            suggestion = issue.get('suggestion', '')
            
            # 提取修正后的文本
            corrected_text = self._extract_corrected_text(suggestion)
            
            # 找到问题文本在哪个段落
            for i, paragraph_text in enumerate(text_content):
                if problem_text in paragraph_text:
                    revisions.append({
                        'paragraph_index': i,
                        'original_text': problem_text,
                        'corrected_text': corrected_text,
                        'reason': f"{issue.get('type', '')} - {issue.get('severity', '')}"
                    })
                    break
        
        # 处理suggestions
        for suggestion in ai_result.suggestions:
            original_text = suggestion.get('original', '')
            suggested_text = suggestion.get('suggested', '')
            
            # 找到原文本在哪个段落
            for i, paragraph_text in enumerate(text_content):
                if original_text in paragraph_text:
                    revisions.append({
                        'paragraph_index': i,
                        'original_text': original_text,
                        'corrected_text': suggested_text,
                        'reason': suggestion.get('reason', '')
                    })
                    break
        
        return revisions
    
    def _extract_corrected_text(self, suggestion: str):
        """从建议中提取修正后的文本"""
        # 尝试从建议中提取修正文本
        if "建议改为：" in suggestion:
            return suggestion.split("建议改为：")[-1].strip()
        elif "应为" in suggestion:
            return suggestion.split("应为")[-1].strip().strip("'\"")
        elif "->" in suggestion:
            return suggestion.split("->")[-1].strip()
        else:
            # 如果无法提取，返回原建议
            return suggestion
    
    def extract_text_content(self, doc: Document):
        """提取文档的文本内容"""
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        return text_content


# 集成到主校对器
class ProofReader:
    """主校对器类 - 支持批注和修订两种模式"""
    
    def __init__(self, api_key: str = None):
        """初始化校对器"""
        self.config = Config()
        if api_key:
            self.config.ai.api_key = api_key
        self.ai_checker = AIChecker(self.config)
        self.console = Console()
        
        self.document_processor = DocumentProcessor()
        self.revisions_proofreader = ProofReaderWithRevisions(api_key)
    
    def proofread_document(self, input_file: str, output_file: str = None, 
                          mode: str = "comments") -> bool:
        """校对文档
        
        Args:
            input_file: 输入文件路径
            output_file: 输出文件路径
            mode: 校对模式 ("comments" 或 "revisions")
        """
        if mode == "revisions":
            return self.revisions_proofreader.proofread_document_with_revisions(input_file, output_file)
        else:
            # 使用原来的批注模式 (这里需要导入原来的实现)
            return self._proofread_with_comments(input_file, output_file)
    
    def _proofread_with_comments(self, input_file: str, output_file: str = None) -> bool:
        """使用批注模式校对（原来的实现）"""
        # 这里调用原来的批注实现
        self.console.print("[yellow]使用批注模式校对（需要导入原始实现）[/yellow]")
        return True
    
    def extract_text_content(self, doc: Document):
        """提取文档的文本内容"""
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        return text_content


# 测试函数
def test_revision_proofreader():
    """测试修订校对器"""
    try:
        # 使用测试API密钥
        api_key = "sk-test"  # 替换为真实的API密钥
        
        proofreader = ProofReaderWithRevisions(api_key)
        
        input_file = "sample_input.docx"
        output_file = "sample_output_revisions.docx"
        
        if os.path.exists(input_file):
            success = proofreader.proofread_document_with_revisions(input_file, output_file)
            if success:
                print(f"✅ 修订校对成功: {output_file}")
            else:
                print("❌ 修订校对失败")
        else:
            print(f"❌ 输入文件不存在: {input_file}")
            
    except Exception as e:
        print(f"❌ 测试失败: {e}")


if __name__ == "__main__":
    test_revision_proofreader() 
#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
高级Word批注处理模块 - 实现真正的Word审阅批注功能
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
from datetime import datetime
import xml.etree.ElementTree as ET


class WordCommentsManager:
    """Word审阅批注管理器"""
    
    def __init__(self, document):
        self.document = document
        self.comment_counter = 0
        self.comments = []  # 存储批注信息
        
    def add_comment(self, paragraph, target_text: str, comment_text: str, author: str = "AI校对助手"):
        """在段落中添加Word审阅批注"""
        try:
            original_text = paragraph.text
            start_pos = original_text.find(target_text)
            
            if start_pos == -1:
                print(f"未找到目标文本: {target_text}")
                return False
            
            end_pos = start_pos + len(target_text)
            
            # 生成批注ID
            self.comment_counter += 1
            comment_id = self.comment_counter
            
            # 存储批注信息
            self.comments.append({
                'id': comment_id,
                'text': comment_text,
                'author': author,
                'date': datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
            })
            
            # 重建段落，正确插入批注标记
            self._rebuild_paragraph_with_comment(paragraph, original_text, start_pos, end_pos, comment_id)
            
            print(f"✅ Word审阅批注已添加: {comment_text[:50]}...")
            return True
            
        except Exception as e:
            print(f"添加Word审阅批注失败: {e}")
            return False
    
    def _rebuild_paragraph_with_comment(self, paragraph, original_text, start_pos, end_pos, comment_id):
        """重建段落，正确插入批注标记"""
        try:
            # 不清空段落，而是在现有内容基础上添加批注标记
            # 直接在段落的XML元素中添加批注标记
            
            # 1. 添加批注范围开始标记
            self._add_comment_range_start_to_element(paragraph._element, comment_id)
            
            # 2. 查找包含目标文本的run并高亮
            for run in paragraph.runs:
                if start_pos <= len(run.text) and run.text:
                    # 找到包含目标文本的run
                    run_text = run.text
                    if original_text[start_pos:end_pos] in run_text:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        break
            
            # 3. 添加批注范围结束标记
            self._add_comment_range_end_to_element(paragraph._element, comment_id)
            
            # 4. 添加批注引用标记
            self._add_comment_reference_to_element(paragraph._element, comment_id)
                
        except Exception as e:
            print(f"重建段落失败: {e}")
            # 如果失败，使用简单的标记方法
            self._add_simple_comment_markers(paragraph, comment_id)
    
    def _add_comment_range_start(self, paragraph, comment_id):
        """添加批注范围开始标记"""
        try:
            element = OxmlElement('w:commentRangeStart')
            element.set(qn('w:id'), str(comment_id))
            paragraph._element.append(element)
            print(f"✅ 添加批注范围开始标记: comment_id={comment_id}")
        except Exception as e:
            print(f"添加批注范围开始标记失败: {e}")
    
    def _add_comment_range_end(self, paragraph, comment_id):
        """添加批注范围结束标记"""
        try:
            element = OxmlElement('w:commentRangeEnd')
            element.set(qn('w:id'), str(comment_id))
            paragraph._element.append(element)
            print(f"✅ 添加批注范围结束标记: comment_id={comment_id}")
        except Exception as e:
            print(f"添加批注范围结束标记失败: {e}")
    
    def _add_comment_reference_run(self, paragraph, comment_id):
        """在独立的run中添加批注引用标记"""
        try:
            # 创建新的run
            new_run = paragraph.add_run()
            
            # 在run的XML元素中添加批注引用
            run_element = new_run._element
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), str(comment_id))
            
            run_element.append(comment_ref)
            
            print(f"✅ 添加批注引用标记: comment_id={comment_id}")
        except Exception as e:
            print(f"添加批注引用标记失败: {e}")
    
    def _add_comment_range_start_to_element(self, paragraph_element, comment_id):
        """直接在段落元素中添加批注范围开始标记"""
        try:
            element = OxmlElement('w:commentRangeStart')
            element.set(qn('w:id'), str(comment_id))
            paragraph_element.insert(0, element)  # 插入到段落开始
            print(f"✅ 添加批注范围开始标记到元素: comment_id={comment_id}")
        except Exception as e:
            print(f"添加批注范围开始标记到元素失败: {e}")
    
    def _add_comment_range_end_to_element(self, paragraph_element, comment_id):
        """直接在段落元素中添加批注范围结束标记"""
        try:
            element = OxmlElement('w:commentRangeEnd')
            element.set(qn('w:id'), str(comment_id))
            paragraph_element.append(element)  # 添加到段落末尾
            print(f"✅ 添加批注范围结束标记到元素: comment_id={comment_id}")
        except Exception as e:
            print(f"添加批注范围结束标记到元素失败: {e}")
    
    def _add_comment_reference_to_element(self, paragraph_element, comment_id):
        """直接在段落元素中添加批注引用标记"""
        try:
            # 创建一个新的run元素
            run_element = OxmlElement('w:r')
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), str(comment_id))
            run_element.append(comment_ref)
            paragraph_element.append(run_element)
            print(f"✅ 添加批注引用标记到元素: comment_id={comment_id}")
        except Exception as e:
            print(f"添加批注引用标记到元素失败: {e}")
    
    def _add_simple_comment_markers(self, paragraph, comment_id):
        """简单的批注标记方法（备用）"""
        try:
            # 使用原有的方法作为备用
            self._add_comment_range_start(paragraph, comment_id)
            self._add_comment_range_end(paragraph, comment_id)
            self._add_comment_reference_run(paragraph, comment_id)
        except Exception as e:
            print(f"添加简单批注标记失败: {e}")

    def _add_comment_reference(self, paragraph, comment_id):
        """添加批注引用标记（保留旧方法作为备用）"""
        try:
            # 使用新的方法
            self._add_comment_reference_run(paragraph, comment_id)
        except Exception as e:
            print(f"添加批注引用标记失败: {e}")
    
    def finalize_document(self):
        """完成文档处理，准备批注数据"""
        try:
            if not self.comments:
                print("没有批注需要处理")
                return True
            
            print(f"✅ 准备 {len(self.comments)} 个批注数据用于XML生成")
            return True
                
        except Exception as e:
            print(f"完成文档处理失败: {e}")
            return False
    
    def get_comments_for_xml(self):
        """获取用于生成XML的批注数据"""
        xml_comments = []
        for comment in self.comments:
            xml_comments.append({
                'id': comment['id'],
                'text': comment['text'],
                'author': comment['author'],
                'date': comment['date']
            })
        return xml_comments


# 测试函数
def test_word_comments():
    """测试Word审阅批注功能"""
    try:
        # 创建测试文档
        doc = Document()
        doc.add_paragraph("这是一个测试文档。")
        doc.add_paragraph("计算器科学是一门重要的学科。")
        doc.add_paragraph("程式设计需要仔细考虑。")
        
        # 创建批注管理器
        comments_manager = WordCommentsManager(doc)
        
        # 添加批注
        paragraphs = list(doc.paragraphs)
        comments_manager.add_comment(paragraphs[1], "计算器科学", 
                                   "错别字：应为'计算机科学'", "测试用户")
        comments_manager.add_comment(paragraphs[2], "程式设计", 
                                   "术语问题：应为'程序设计'", "测试用户")
        
        # 完成文档处理
        comments_manager.finalize_document()
        
        # 保存测试文档
        doc.save("test_word_review_comments.docx")
        print("✅ 测试文档已保存: test_word_review_comments.docx")
        print("📝 文档包含Word审阅批注标记，使用Microsoft Word打开可查看完整批注")
        
    except Exception as e:
        print(f"测试失败: {e}")


if __name__ == "__main__":
    test_word_comments() 
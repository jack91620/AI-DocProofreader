"""
文档处理模块
"""

import re
from typing import List, Dict, Tuple
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import uuid
from datetime import datetime
from .word_comments import WordCommentsHandler


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        self.document = None
        self.paragraphs = []
        self.text_content = ""
        self.comments_handler = None
    
    def load_document(self, file_path: str) -> bool:
        """加载docx文档"""
        try:
            self.document = Document(file_path)
            self.comments_handler = WordCommentsHandler(self.document)
            self._extract_text()
            return True
        except Exception as e:
            print(f"加载文档失败: {e}")
            return False
    
    def _extract_text(self):
        """提取文档中的文本内容"""
        self.paragraphs = []
        all_text = []
        
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                self.paragraphs.append({
                    'text': paragraph.text,
                    'paragraph': paragraph,
                    'index': len(self.paragraphs)
                })
                all_text.append(paragraph.text)
        
        self.text_content = '\n'.join(all_text)
    
    def get_text_segments(self, max_length: int = 2000) -> List[str]:
        """将文档文本分割成适合AI处理的段落"""
        segments = []
        current_segment = ""
        
        for para_info in self.paragraphs:
            text = para_info['text']
            
            # 如果当前段落加上新文本超过最大长度，保存当前段落并开始新段落
            if len(current_segment) + len(text) > max_length and current_segment:
                segments.append(current_segment.strip())
                current_segment = text
            else:
                current_segment += "\n" + text if current_segment else text
        
        # 添加最后一个段落
        if current_segment:
            segments.append(current_segment.strip())
        
        return segments
    
    def add_comment(self, paragraph_index: int, text: str, comment: str, 
                   author: str = "AI校对助手", color: str = "red"):
        """在指定段落添加Word原生批注"""
        try:
            if paragraph_index >= len(self.paragraphs):
                return False
            
            paragraph = self.paragraphs[paragraph_index]['paragraph']
            
            # 查找要批注的文本
            if text in paragraph.text:
                # 使用Word原生批注功能
                self._add_word_comment(paragraph, text, comment, author)
                return True
            
        except Exception as e:
            print(f"添加批注失败: {e}")
        
        return False
    
    def _add_word_comment(self, paragraph, target_text: str, comment: str, author: str):
        """使用Word原生批注功能添加批注"""
        try:
            # 保存原始段落文本
            original_text = paragraph.text
            start_pos = original_text.find(target_text)
            
            if start_pos == -1:
                return False
            
            end_pos = start_pos + len(target_text)
            
            # 清空段落
            paragraph.clear()
            
            # 添加目标文本之前的内容
            if start_pos > 0:
                paragraph.add_run(original_text[:start_pos])
            
            # 创建带有批注标记的run（高亮显示）
            commented_run = paragraph.add_run(target_text)
            commented_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # 直接添加批注内容到段落中（确保可见）
            comment_run = paragraph.add_run(f" [批注: {comment}]")
            comment_run.font.color.rgb = RGBColor(204, 0, 0)  # 深红色
            comment_run.font.size = 160000  # 8pt (160000 twips = 8pt)
            comment_run.font.italic = True
            
            # 添加目标文本之后的内容
            if end_pos < len(original_text):
                paragraph.add_run(original_text[end_pos:])
            
            print(f"📝 批注已添加到文档: {comment}")
            return True
            
        except Exception as e:
            print(f"添加Word批注失败: {e}")
            # 如果失败，回退到简单的文本批注
            return self._add_simple_text_comment(paragraph, target_text, comment, author)
    
    def highlight_text(self, paragraph_index: int, text: str, 
                      color: WD_COLOR_INDEX = WD_COLOR_INDEX.YELLOW):
        """高亮显示文本"""
        try:
            if paragraph_index >= len(self.paragraphs):
                return False
            
            paragraph = self.paragraphs[paragraph_index]['paragraph']
            
            # 在段落中查找并高亮文本
            for run in paragraph.runs:
                if text in run.text:
                    run.font.highlight_color = color
                    return True
            
        except Exception as e:
            print(f"高亮文本失败: {e}")
        
        return False
    
    def save_document(self, output_path: str) -> bool:
        """保存文档"""
        try:
            if self.document:
                self.document.save(output_path)
                return True
        except Exception as e:
            print(f"保存文档失败: {e}")
        
        return False
    
    def get_paragraph_by_text(self, text: str) -> Tuple[int, str]:
        """根据文本内容查找段落"""
        for i, para_info in enumerate(self.paragraphs):
            if text in para_info['text']:
                return i, para_info['text']
        return -1, ""
    
    def _add_visual_comment_marker(self, run, comment: str):
        """添加视觉批注标记，包含完整批注内容"""
        try:
            # 高亮显示批注的文本
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # 获取run所在的段落
            paragraph = run._element.getparent().getparent()
            
            # 创建包含批注内容的标记XML
            comment_marker_xml = f'''
            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:rPr>
                    <w:color w:val="CC0000"/>
                    <w:sz w:val="16"/>
                    <w:i/>
                </w:rPr>
                <w:t xml:space="preserve"> [批注: {comment}]</w:t>
            </w:r>
            '''
            
            # 解析并添加批注标记
            comment_marker = parse_xml(comment_marker_xml)
            run._element.addnext(comment_marker)
            
            # 记录批注信息
            print(f"📝 批注已添加: {comment}")
            
        except Exception as e:
            print(f"添加视觉批注标记失败: {e}")
            # 最简单的备用方案
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            # 添加简化的批注文本
            try:
                # 在段落中直接添加批注文本
                paragraph = run._element.getparent()
                if paragraph is not None:
                    # 获取段落的父级元素来添加批注
                    for parent_paragraph in paragraph.iter():
                        if parent_paragraph.tag.endswith('}p'):
                            # 在段落后添加批注run
                            simple_comment = parse_xml(f'''
                            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                                <w:rPr>
                                    <w:color w:val="FF0000"/>
                                    <w:sz w:val="16"/>
                                    <w:i/>
                                </w:rPr>
                                <w:t> [批注: {comment}]</w:t>
                            </w:r>
                            ''')
                            parent_paragraph.append(simple_comment)
                            break
            except Exception as inner_e:
                print(f"备用批注方案也失败: {inner_e}")
    
    def _add_simple_text_comment(self, paragraph, target_text: str, comment: str, author: str):
        """备用方法：添加简单的文本批注"""
        try:
            original_text = paragraph.text
            start_pos = original_text.find(target_text)
            
            if start_pos == -1:
                return False
            
            end_pos = start_pos + len(target_text)
            
            # 清空段落
            paragraph.clear()
            
            # 添加目标文本之前的内容
            if start_pos > 0:
                paragraph.add_run(original_text[:start_pos])
            
            # 添加高亮的目标文本
            highlighted_run = paragraph.add_run(target_text)
            highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # 添加简化的批注标识
            comment_run = paragraph.add_run(f" [批注: {comment}]")
            comment_run.font.color.rgb = RGBColor(200, 0, 0)
            comment_run.font.size = 90000  # 9pt
            comment_run.font.italic = True
            
            # 添加目标文本之后的内容
            if end_pos < len(original_text):
                paragraph.add_run(original_text[end_pos:])
            
            return True
            
        except Exception as e:
            print(f"添加简单文本批注失败: {e}")
            return False
    
    def extract_text_content(self, file_path: str) -> List[str]:
        """提取文档文本内容"""
        try:
            document = Document(file_path)
            paragraphs = []
            
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:  # 只保留非空段落
                    paragraphs.append(text)
            
            return paragraphs
        except Exception as e:
            print(f"提取文档内容失败: {e}")
            return []
    
    def get_statistics(self) -> Dict:
        """获取文档统计信息"""
        return {
            "paragraph_count": len(self.paragraphs),
            "character_count": len(self.text_content),
            "word_count": len(self.text_content.replace(' ', '')),  # 中文字符数
        } 
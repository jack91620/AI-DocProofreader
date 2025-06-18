#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
真正的Word跟踪更改功能模块
实现Microsoft Word审阅中的修订功能
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import zipfile
import tempfile
import os
import xml.etree.ElementTree as ET


class WordTrackChangesManager:
    """Word跟踪更改管理器 - 生成真正的Word修订标记"""
    
    def __init__(self, document):
        self.document = document
        self.revision_counter = 0
        self.author = "AI校对助手"
        self.date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        self.revisions_data = []
        self.paragraph_changes = {}  # 跟踪每个段落的修改
    
    def add_tracked_change(self, paragraph, original_text: str, corrected_text: str, reason: str = ""):
        """添加跟踪更改（真正的Word修订）"""
        try:
            # 生成修订ID
            self.revision_counter += 1
            del_revision_id = self.revision_counter
            ins_revision_id = self.revision_counter + 1000
            
            # 存储修订数据
            self.revisions_data.append({
                'paragraph': paragraph,
                'original_text': original_text,
                'corrected_text': corrected_text,
                'del_id': del_revision_id,
                'ins_id': ins_revision_id,
                'author': self.author,
                'date': self.date,
                'reason': reason
            })
            
            # 跟踪段落修改，批量处理
            para_id = id(paragraph)
            if para_id not in self.paragraph_changes:
                self.paragraph_changes[para_id] = {
                    'paragraph': paragraph,
                    'original_text': paragraph.text,
                    'changes': []
                }
            
            self.paragraph_changes[para_id]['changes'].append({
                'original_text': original_text,
                'corrected_text': corrected_text,
                'del_id': del_revision_id,
                'ins_id': ins_revision_id,
                'reason': reason
            })
            
            print(f"✅ 已添加跟踪更改: {original_text} -> {corrected_text}")
            return True
            
        except Exception as e:
            print(f"添加跟踪更改失败: {e}")
            return False
    
    def apply_all_changes(self):
        """应用所有跟踪更改到文档"""
        for para_id, para_data in self.paragraph_changes.items():
            paragraph = para_data['paragraph']
            original_full_text = para_data['original_text']
            changes = para_data['changes']
            
            # 应用所有更改到段落
            self._apply_changes_to_paragraph(paragraph, original_full_text, changes)
    
    def _apply_changes_to_paragraph(self, paragraph, original_text, changes):
        """将多个更改应用到单个段落（简化版本）"""
        try:
            # 按照在原文中的位置排序（从后往前，避免位置变化）
            changes_with_pos = []
            for change in changes:
                pos = original_text.find(change['original_text'])
                if pos != -1:
                    changes_with_pos.append((pos, change))
            
            # 按位置降序排序
            changes_with_pos.sort(key=lambda x: x[0], reverse=True)
            
            # 清空段落
            paragraph.clear()
            
            # 如果没有更改，直接添加原文
            if not changes_with_pos:
                paragraph.add_run(original_text)
                return
            
            # 构建新的段落内容
            current_text = original_text
            
            # 从后往前处理每个更改
            for pos, change in changes_with_pos:
                original_part = change['original_text']
                corrected_part = change['corrected_text']
                del_id = change['del_id']
                ins_id = change['ins_id']
                
                # 分割文本
                before_part = current_text[:pos]
                after_part = current_text[pos + len(original_part):]
                
                # 重新组合文本
                current_text = before_part + "@@DEL:" + str(del_id) + ":" + original_part + "@@" + "@@INS:" + str(ins_id) + ":" + corrected_part + "@@" + after_part
            
            # 解析并构建段落
            self._build_paragraph_from_marked_text(paragraph, current_text)
                    
        except Exception as e:
            print(f"应用段落更改失败: {e}")
            # 如果失败，至少保留原始文本
            paragraph.clear()
            paragraph.add_run(original_text)
    
    def _build_paragraph_from_marked_text(self, paragraph, marked_text):
        """从标记文本构建段落"""
        try:
            # 解析标记文本
            pos = 0
            while pos < len(marked_text):
                # 查找下一个删除标记
                del_start = marked_text.find("@@DEL:", pos)
                if del_start == -1:
                    # 没有更多删除标记，添加剩余文本
                    remaining = marked_text[pos:]
                    if remaining:
                        paragraph.add_run(remaining)
                    break
                
                # 添加删除标记之前的文本
                if del_start > pos:
                    before_text = marked_text[pos:del_start]
                    paragraph.add_run(before_text)
                
                # 解析删除标记
                del_end = marked_text.find("@@", del_start + 6)
                if del_end == -1:
                    break
                
                del_content = marked_text[del_start + 6:del_end]
                parts = del_content.split(":", 1)
                if len(parts) == 2:
                    del_id = parts[0]
                    del_text = parts[1]
                    
                    # 创建删除元素
                    del_element = self._create_deletion_element(del_text, del_id)
                    paragraph._element.append(del_element)
                
                # 查找对应的插入标记
                ins_start = marked_text.find("@@INS:", del_end)
                if ins_start == del_end + 2:  # 紧接着的插入标记
                    ins_end = marked_text.find("@@", ins_start + 6)
                    if ins_end != -1:
                        ins_content = marked_text[ins_start + 6:ins_end]
                        parts = ins_content.split(":", 1)
                        if len(parts) == 2:
                            ins_id = parts[0]
                            ins_text = parts[1]
                            
                            # 创建插入元素
                            ins_element = self._create_insertion_element(ins_text, ins_id)
                            paragraph._element.append(ins_element)
                        
                        pos = ins_end + 2
                    else:
                        pos = del_end + 2
                else:
                    pos = del_end + 2
                    
        except Exception as e:
            print(f"构建段落失败: {e}")
            # 如果失败，清理并添加原始文本
            paragraph.clear()
            clean_text = marked_text.replace("@@DEL:", "").replace("@@INS:", "").replace("@@", "")
            paragraph.add_run(clean_text)
    
    def _create_deletion_element(self, text, revision_id):
        """创建删除元素 (w:del)"""
        # 创建删除元素
        del_element = OxmlElement('w:del')
        del_element.set(qn('w:id'), str(revision_id))
        del_element.set(qn('w:author'), self.author)
        del_element.set(qn('w:date'), self.date)
        
        # 创建删除的run
        del_run = OxmlElement('w:r')
        
        # 添加删除文本
        del_text = OxmlElement('w:delText')
        del_text.text = text
        del_run.append(del_text)
        
        del_element.append(del_run)
        return del_element
    
    def _create_insertion_element(self, text, revision_id):
        """创建插入元素 (w:ins)"""
        # 创建插入元素
        ins_element = OxmlElement('w:ins')
        ins_element.set(qn('w:id'), str(revision_id))
        ins_element.set(qn('w:author'), self.author)
        ins_element.set(qn('w:date'), self.date)
        
        # 创建插入的run
        ins_run = OxmlElement('w:r')
        
        # 添加插入文本
        ins_text = OxmlElement('w:t')
        ins_text.text = text
        ins_run.append(ins_text)
        
        ins_element.append(ins_run)
        return ins_element


def enable_track_changes_in_docx(docx_path, output_path, revisions_data):
    """在Word文档中启用跟踪更改并添加修订"""
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            # 解压docx文件
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # 修改settings.xml以启用跟踪更改
            enable_track_changes_setting(temp_dir)
            
            # 验证document.xml中的修订标记
            verify_document_revisions(temp_dir)
            
            # 重新打包为docx文件
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arc_name = os.path.relpath(file_path, temp_dir)
                        zip_ref.write(file_path, arc_name)
            
            print(f"✅ 成功启用Word跟踪更改: {output_path}")
            return True
            
    except Exception as e:
        print(f"❌ 启用跟踪更改失败: {e}")
        return False


def enable_track_changes_setting(temp_dir):
    """在settings.xml中启用跟踪更改"""
    try:
        settings_path = os.path.join(temp_dir, 'word', 'settings.xml')
        
        if os.path.exists(settings_path):
            # 解析现有的settings.xml
            tree = ET.parse(settings_path)
            root = tree.getroot()
        else:
            # 创建新的settings.xml
            ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            root = ET.Element(f'{{{ns_w}}}settings')
            tree = ET.ElementTree(root)
        
        # 定义命名空间
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # 检查是否已存在trackRevisions设置
        track_revisions = root.find('.//w:trackRevisions', ns)
        
        if track_revisions is None:
            # 添加trackRevisions设置
            track_revisions = ET.SubElement(root, f"{{{ns['w']}}}trackRevisions")
        
        # 确保跟踪更改被启用
        track_revisions.set(f"{{{ns['w']}}}val", "1")
        
        # 保存settings.xml
        tree.write(settings_path, encoding='utf-8', xml_declaration=True)
        print("✅ 已在settings.xml中启用跟踪更改")
        
    except Exception as e:
        print(f"❌ 启用跟踪更改设置失败: {e}")


def verify_document_revisions(temp_dir):
    """验证document.xml中的修订标记"""
    try:
        document_path = os.path.join(temp_dir, 'word', 'document.xml')
        
        if os.path.exists(document_path):
            with open(document_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 统计修订标记
            del_count = content.count('<w:del ')
            ins_count = content.count('<w:ins ')
            deltext_count = content.count('<w:delText>')
            
            print(f"📊 document.xml修订标记统计:")
            print(f"   - w:del (删除标记): {del_count}")
            print(f"   - w:ins (插入标记): {ins_count}")
            print(f"   - w:delText (删除文本): {deltext_count}")
            
            if del_count > 0 or ins_count > 0:
                print("✅ 发现Word修订标记")
            else:
                print("⚠️  未发现Word修订标记")
        
    except Exception as e:
        print(f"❌ 验证修订标记失败: {e}")


# 测试函数
def test_word_track_changes():
    """测试Word跟踪更改功能"""
    try:
        # 创建测试文档
        doc = Document()
        doc.add_paragraph("这是一个测试文档。")
        doc.add_paragraph("计算器科学是一门重要的学科。")
        doc.add_paragraph("程式设计需要仔细考虑。")
        
        # 创建跟踪更改管理器
        track_changes_manager = WordTrackChangesManager(doc)
        
        # 添加跟踪更改
        paragraphs = list(doc.paragraphs)
        track_changes_manager.add_tracked_change(paragraphs[1], "计算器科学", "计算机科学", "错别字修正")
        track_changes_manager.add_tracked_change(paragraphs[2], "程式设计", "程序设计", "术语统一")
        
        # 应用所有更改
        track_changes_manager.apply_all_changes()
        
        # 保存临时文档
        temp_file = "test_word_track_changes_temp.docx"
        doc.save(temp_file)
        
        # 启用跟踪更改并生成最终文档
        output_file = "test_word_track_changes.docx"
        if enable_track_changes_in_docx(temp_file, output_file, track_changes_manager.revisions_data):
            print(f"✅ Word跟踪更改测试文档已创建: {output_file}")
            print("📝 现在可以在Microsoft Word中查看真正的修订功能")
            
            # 清理临时文件
            if os.path.exists(temp_file):
                os.remove(temp_file)
        else:
            print("❌ 创建失败")
        
    except Exception as e:
        print(f"❌ 测试失败: {e}")


if __name__ == "__main__":
    test_word_track_changes() 
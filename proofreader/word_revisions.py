#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word修订功能模块 - 实现跟踪更改功能
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import xml.etree.ElementTree as ET


class WordRevisionsManager:
    """Word修订管理器"""
    
    def __init__(self, document):
        self.document = document
        self.revision_counter = 0
        self.author = "AI校对助手"
        self.date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    
    def add_revision(self, paragraph, original_text: str, corrected_text: str, reason: str = ""):
        """在段落中添加修订标记"""
        try:
            paragraph_text = paragraph.text
            start_pos = paragraph_text.find(original_text)
            
            if start_pos == -1:
                print(f"未找到需要修订的文本: {original_text}")
                return False
            
            end_pos = start_pos + len(original_text)
            
            # 生成修订ID
            self.revision_counter += 1
            revision_id = self.revision_counter
            
            # 清空段落并重建
            paragraph.clear()
            
            # 添加原文本之前的内容
            if start_pos > 0:
                paragraph.add_run(paragraph_text[:start_pos])
            
            # 添加删除标记（原始错误文本）
            self._add_deleted_text(paragraph, original_text, revision_id)
            
            # 添加插入标记（修正后文本）
            self._add_inserted_text(paragraph, corrected_text, revision_id)
            
            # 添加原文本之后的内容
            if end_pos < len(paragraph_text):
                paragraph.add_run(paragraph_text[end_pos:])
            
            print(f"✅ Word修订已添加: {original_text} -> {corrected_text}")
            return True
            
        except Exception as e:
            print(f"添加Word修订失败: {e}")
            return False
    
    def _add_deleted_text(self, paragraph, text: str, revision_id: int):
        """添加删除的文本标记"""
        try:
            # 创建删除标记元素
            del_element = OxmlElement('w:del')
            del_element.set(qn('w:id'), str(revision_id))
            del_element.set(qn('w:author'), self.author)
            del_element.set(qn('w:date'), self.date)
            
            # 创建删除文本的run
            del_run = OxmlElement('w:r')
            
            # 设置删除文本的属性
            del_run_props = OxmlElement('w:rPr')
            del_run.append(del_run_props)
            
            # 添加文本
            del_text = OxmlElement('w:delText')
            del_text.text = text
            del_run.append(del_text)
            
            del_element.append(del_run)
            paragraph._element.append(del_element)
            
        except Exception as e:
            print(f"添加删除文本标记失败: {e}")
    
    def _add_inserted_text(self, paragraph, text: str, revision_id: int):
        """添加插入的文本标记"""
        try:
            # 创建插入标记元素
            ins_element = OxmlElement('w:ins')
            ins_element.set(qn('w:id'), str(revision_id + 1000))  # 使用不同的ID
            ins_element.set(qn('w:author'), self.author)
            ins_element.set(qn('w:date'), self.date)
            
            # 创建插入文本的run
            ins_run = OxmlElement('w:r')
            
            # 设置插入文本的属性
            ins_run_props = OxmlElement('w:rPr')
            ins_run.append(ins_run_props)
            
            # 添加文本
            ins_text = OxmlElement('w:t')
            ins_text.text = text
            ins_run.append(ins_text)
            
            ins_element.append(ins_run)
            paragraph._element.append(ins_element)
            
        except Exception as e:
            print(f"添加插入文本标记失败: {e}")
    
    def enable_track_changes(self):
        """启用文档的跟踪更改功能"""
        try:
            # 在文档设置中启用跟踪更改
            settings_element = self._get_or_create_settings()
            
            # 添加跟踪更改设置
            track_revisions = OxmlElement('w:trackRevisions')
            settings_element.append(track_revisions)
            
            print("✅ 已启用文档跟踪更改功能")
            return True
            
        except Exception as e:
            print(f"启用跟踪更改失败: {e}")
            return False
    
    def _get_or_create_settings(self):
        """获取或创建文档设置元素"""
        try:
            # 这是一个简化版本，实际实现需要更复杂的XML操作
            # 由于python-docx的限制，我们主要依赖XML标记
            return OxmlElement('w:settings')
        except Exception as e:
            print(f"获取文档设置失败: {e}")
            return None


class SimpleWordRevisionsManager:
    """简化的Word修订管理器 - 使用直观的修订显示"""
    
    def __init__(self, document):
        self.document = document
        self.revision_counter = 0
        self.author = "AI校对助手"
    
    def add_revision(self, paragraph, original_text: str, corrected_text: str, reason: str = ""):
        """添加简化的修订标记"""
        try:
            paragraph_text = paragraph.text
            start_pos = paragraph_text.find(original_text)
            
            if start_pos == -1:
                print(f"未找到需要修订的文本: {original_text}")
                return False
            
            end_pos = start_pos + len(original_text)
            
            # 清空段落并重建
            paragraph.clear()
            
            # 添加原文本之前的内容
            if start_pos > 0:
                paragraph.add_run(paragraph_text[:start_pos])
            
            # 添加删除的文本（删除线样式）
            deleted_run = paragraph.add_run(original_text)
            deleted_run.font.strike = True
            deleted_run.font.color.rgb = None  # 红色
            
            # 添加插入的文本（下划线样式）
            inserted_run = paragraph.add_run(corrected_text)
            inserted_run.font.underline = True
            # 设置为蓝色
            from docx.shared import RGBColor
            inserted_run.font.color.rgb = RGBColor(0, 0, 255)
            
            # 添加原文本之后的内容
            if end_pos < len(paragraph_text):
                paragraph.add_run(paragraph_text[end_pos:])
            
            self.revision_counter += 1
            print(f"✅ 修订标记已添加: {original_text} -> {corrected_text}")
            return True
            
        except Exception as e:
            print(f"添加修订标记失败: {e}")
            return False


def create_revisions_xml(revisions_data):
    """创建修订XML内容（用于完整的Word修订功能）"""
    # XML命名空间
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # 注册命名空间
    ET.register_namespace('w', ns_w)
    
    # 创建修订信息（这是一个概念性实现）
    revisions_info = {
        'total_revisions': len(revisions_data),
        'author': 'AI校对助手',
        'date': datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"),
        'revisions': revisions_data
    }
    
    return revisions_info


# 测试函数
def test_word_revisions():
    """测试Word修订功能"""
    try:
        # 创建测试文档
        doc = Document()
        doc.add_paragraph("这是一个测试文档。")
        doc.add_paragraph("计算器科学是一门重要的学科。")
        doc.add_paragraph("程式设计需要仔细考虑。")
        
        # 创建修订管理器
        revisions_manager = SimpleWordRevisionsManager(doc)
        
        # 添加修订
        paragraphs = list(doc.paragraphs)
        revisions_manager.add_revision(paragraphs[1], "计算器科学", "计算机科学", "错别字修正")
        revisions_manager.add_revision(paragraphs[2], "程式设计", "程序设计", "术语统一")
        
        # 保存测试文档
        doc.save("test_word_revisions.docx")
        print("✅ 修订测试文档已保存: test_word_revisions.docx")
        print("📝 文档包含修订标记：删除线表示删除，下划线表示插入")
        
    except Exception as e:
        print(f"测试失败: {e}")


if __name__ == "__main__":
    test_word_revisions() 
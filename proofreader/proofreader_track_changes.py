#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
使用真正Word跟踪更改功能的校对器
"""

import os
import sys
from typing import Optional
from rich.console import Console
from docx import Document
from datetime import datetime

from .config import Config
from .document import DocumentProcessor
from .ai_checker import AIChecker, ProofreadingResult
from .word_track_changes import WordTrackChangesManager, enable_track_changes_in_docx


class ProofReaderWithTrackChanges:
    """使用真正Word跟踪更改功能的校对器"""
    
    def __init__(self, api_key: str = None):
        """初始化校对器"""
        self.config = Config()
        if api_key:
            self.config.ai.api_key = api_key
        self.ai_checker = AIChecker(self.config)
        self.console = Console()
        
        self.document_processor = DocumentProcessor()
    
    def proofread_with_track_changes(self, input_file: str, output_file: str = None) -> bool:
        """使用Word跟踪更改功能进行校对"""
        try:
            # 生成输出文件名
            if not output_file:
                output_file = input_file.replace('.docx', '_tracked.docx')
            
            self.console.print(f"[green]开始Word跟踪更改校对：{input_file}[/green]")
            
            # 读取文档
            doc = Document(input_file)
            
            # 创建跟踪更改管理器
            track_changes_manager = WordTrackChangesManager(doc)
            
            # 提取文本内容
            text_content = self.extract_text_content(doc)
            self.console.print(f"[blue]提取文本内容: {len(text_content)} 个段落[/blue]")
            
            # 进行AI校对
            self.console.print("[bold]开始AI校对...")
            ai_result = self.ai_checker.check_text(' '.join(text_content))
            
            # 转换AI校对结果为跟踪更改格式
            changes = self._convert_ai_result_to_track_changes(ai_result, text_content)
            self.console.print(f"[green]✅ AI校对完成，发现 {len(changes)} 个需要跟踪更改的问题[/green]")
            
            # 应用跟踪更改
            change_count = self.apply_track_changes(doc, changes, track_changes_manager)
            
            # 应用所有跟踪更改到文档
            self.console.print("[blue]正在应用所有跟踪更改到文档...[/blue]")
            track_changes_manager.apply_all_changes()
            
            # 保存临时文档
            temp_file = output_file.replace('.docx', '_temp.docx')
            doc.save(temp_file)
            
            # 启用Word跟踪更改并生成最终文档
            if enable_track_changes_in_docx(temp_file, output_file, track_changes_manager.revisions_data):
                # 清理临时文件
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    
                self.console.print(f"[green]✅ Word跟踪更改校对完成，输出文件: {output_file}[/green]")
                self.console.print(f"[blue]📝 已应用 {change_count} 个跟踪更改，现在可以在Word审阅功能中查看[/blue]")
                return True
            else:
                self.console.print(f"[red]❌ 启用跟踪更改失败[/red]")
                return False
            
        except Exception as e:
            self.console.print(f"[red]❌ Word跟踪更改校对失败: {e}[/red]")
            return False
    
    def apply_track_changes(self, doc: Document, changes: list, track_changes_manager: WordTrackChangesManager):
        """应用跟踪更改到文档"""
        change_count = 0
        
        for change in changes:
            paragraph_index = change.get('paragraph_index', 0)
            original_text = change.get('original_text', '')
            corrected_text = change.get('corrected_text', '')
            reason = change.get('reason', '')
            
            # 获取对应段落
            if paragraph_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[paragraph_index]
                
                # 应用跟踪更改
                if track_changes_manager.add_tracked_change(paragraph, original_text, corrected_text, reason):
                    change_count += 1
                    self.console.print(f"[green]✅ 跟踪更改 {change_count}: {original_text} -> {corrected_text}[/green]")
                else:
                    self.console.print(f"[red]❌ 跟踪更改失败: {original_text}[/red]")
        
        return change_count
    
    def _convert_ai_result_to_track_changes(self, ai_result: ProofreadingResult, text_content: list):
        """将AI校对结果转换为跟踪更改格式"""
        changes = []
        
        # 处理issues
        for issue in ai_result.issues:
            problem_text = issue.get('text', '')
            suggestion = issue.get('suggestion', '')
            
            # 提取修正后的文本
            corrected_text = self._extract_corrected_text(suggestion)
            
            # 如果修正文本与原文本不同，才添加跟踪更改
            if corrected_text and corrected_text != problem_text:
                # 找到问题文本在哪个段落
                for i, paragraph_text in enumerate(text_content):
                    if problem_text in paragraph_text:
                        changes.append({
                            'paragraph_index': i,
                            'original_text': problem_text,
                            'corrected_text': corrected_text,
                            'reason': f"{issue.get('type', '')} - {issue.get('severity', '')}"
                        })
                        break
        
        # 处理suggestions
        for suggestion in ai_result.suggestions:
            original_text = suggestion.get('original', '')
            suggested_text = suggestion.get('suggested', '')
            
            # 如果建议文本与原文本不同，才添加跟踪更改
            if suggested_text and suggested_text != original_text:
                # 找到原文本在哪个段落
                for i, paragraph_text in enumerate(text_content):
                    if original_text in paragraph_text:
                        changes.append({
                            'paragraph_index': i,
                            'original_text': original_text,
                            'corrected_text': suggested_text,
                            'reason': suggestion.get('reason', '')
                        })
                        break
        
        return changes
    
    def _extract_corrected_text(self, suggestion: str):
        """从建议中提取修正后的文本"""
        # 尝试从建议中提取修正文本
        if "建议改为：" in suggestion:
            return suggestion.split("建议改为：")[-1].strip()
        elif "应为" in suggestion:
            return suggestion.split("应为")[-1].strip().strip("'\"")
        elif "->" in suggestion:
            return suggestion.split("->")[-1].strip()
        elif "改为" in suggestion:
            return suggestion.split("改为")[-1].strip().strip("'\"")
        else:
            # 如果无法提取，返回空字符串
            return ""
    
    def extract_text_content(self, doc: Document):
        """提取文档的文本内容"""
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        return text_content


# 测试函数
def test_track_changes_proofreader():
    """测试跟踪更改校对器"""
    try:
        # 使用测试API密钥
        api_key = "sk-test"  # 替换为真实的API密钥
        
        proofreader = ProofReaderWithTrackChanges(api_key)
        
        input_file = "sample_input.docx"
        output_file = "sample_output_track_changes.docx"
        
        if os.path.exists(input_file):
            success = proofreader.proofread_with_track_changes(input_file, output_file)
            if success:
                print(f"✅ Word跟踪更改校对成功: {output_file}")
            else:
                print("❌ Word跟踪更改校对失败")
        else:
            print(f"❌ 输入文件不存在: {input_file}")
            
    except Exception as e:
        print(f"❌ 测试失败: {e}")


if __name__ == "__main__":
    test_track_changes_proofreader() 
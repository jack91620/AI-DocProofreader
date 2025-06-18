#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
验证Word文档修订功能
"""

import zipfile
import xml.etree.ElementTree as ET
import os
from docx import Document


def verify_word_revisions(docx_file):
    """验证Word文档的修订标记"""
    print(f"🔍 验证修订文档: {docx_file}")
    
    if not os.path.exists(docx_file):
        print(f"❌ 文件不存在: {docx_file}")
        return False
    
    try:
        # 使用python-docx检查文档
        doc = Document(docx_file)
        
        # 统计段落和run
        total_paragraphs = len(doc.paragraphs)
        total_runs = 0
        strike_through_runs = 0
        underlined_runs = 0
        colored_runs = 0
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                total_runs += 1
                
                # 检查删除线（表示删除的文本）
                if run.font.strike:
                    strike_through_runs += 1
                
                # 检查下划线（表示插入的文本）
                if run.font.underline:
                    underlined_runs += 1
                
                # 检查颜色
                if run.font.color.rgb is not None:
                    colored_runs += 1
        
        print(f"📊 文档统计:")
        print(f"   - 总段落数: {total_paragraphs}")
        print(f"   - 总run数: {total_runs}")
        print(f"   - 删除线标记 (删除): {strike_through_runs}")
        print(f"   - 下划线标记 (插入): {underlined_runs}")
        print(f"   - 彩色文本: {colored_runs}")
        
        # 检查XML结构
        check_revision_xml_structure(docx_file)
        
        print(f"✅ 修订验证完成!")
        return True
        
    except Exception as e:
        print(f"❌ 验证失败: {e}")
        return False


def check_revision_xml_structure(docx_file):
    """检查Word文档的XML修订结构"""
    try:
        with zipfile.ZipFile(docx_file, 'r') as zip_ref:
            file_list = zip_ref.namelist()
            
            # 检查document.xml
            if 'word/document.xml' in file_list:
                with zip_ref.open('word/document.xml') as f:
                    document_content = f.read().decode('utf-8')
                
                # 统计修订相关的XML元素
                del_count = document_content.count('<w:del ')
                ins_count = document_content.count('<w:ins ')
                deltext_count = document_content.count('<w:delText>')
                
                print(f"🎯 XML修订标记:")
                print(f"   - w:del (删除标记): {del_count}")
                print(f"   - w:ins (插入标记): {ins_count}")
                print(f"   - w:delText (删除文本): {deltext_count}")
                
                # 检查settings.xml
                if 'word/settings.xml' in file_list:
                    with zip_ref.open('word/settings.xml') as f:
                        settings_content = f.read().decode('utf-8')
                    
                    track_revisions = 'trackRevisions' in settings_content
                    print(f"⚙️  跟踪更改设置: {'✅ 已启用' if track_revisions else '❌ 未启用'}")
                else:
                    print("⚙️  settings.xml: ❌ 文件不存在")
            
    except Exception as e:
        print(f"❌ XML结构检查失败: {e}")


def show_revision_content(docx_file):
    """显示修订文档的内容示例"""
    try:
        doc = Document(docx_file)
        print(f"\n📖 修订内容示例 ({docx_file}):")
        
        revision_count = 0
        for i, paragraph in enumerate(doc.paragraphs):
            has_revisions = False
            paragraph_text = ""
            
            for run in paragraph.runs:
                if run.font.strike:
                    # 删除的文本
                    paragraph_text += f"[删除: {run.text}]"
                    has_revisions = True
                elif run.font.underline and run.font.color.rgb is not None:
                    # 插入的文本
                    paragraph_text += f"[插入: {run.text}]"
                    has_revisions = True
                else:
                    paragraph_text += run.text
            
            if has_revisions:
                revision_count += 1
                print(f"   修订 {revision_count}: {paragraph_text[:100]}...")
                
                if revision_count >= 3:  # 只显示前3个修订
                    break
        
        if revision_count == 0:
            print("   未发现明显的修订标记")
        else:
            print(f"   (显示前3个，共发现修订段落数量未完全统计)")
            
    except Exception as e:
        print(f"❌ 显示修订内容失败: {e}")


if __name__ == "__main__":
    # 验证修订文档
    revision_files = [
        "sample_output_revisions.docx",
        "sample_output_revisions2.docx",
        "test_word_revisions.docx"
    ]
    
    print("🔍 Word修订功能验证报告")
    print("=" * 50)
    
    for revision_file in revision_files:
        if os.path.exists(revision_file):
            print(f"\n📝 验证文件: {revision_file}")
            verify_word_revisions(revision_file)
            show_revision_content(revision_file)
            print("-" * 40)
        else:
            print(f"\n❌ 文件不存在: {revision_file}")
    
    print("\n✅ 验证完成!")
    print("\n💡 使用说明:")
    print("   1. 删除线文本表示需要删除的原始内容")
    print("   2. 蓝色下划线文本表示新插入的修正内容")
    print("   3. 在Microsoft Word中可以使用审阅功能查看和管理这些修订")
    print("   4. 可以逐个接受或拒绝修订建议") 